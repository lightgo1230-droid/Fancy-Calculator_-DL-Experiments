"""
애니메이션 영화 데이터셋 - PyTorch 머신러닝 3종 프로젝트 분석 결과 Word 보고서 생성
"""
from docx import Document
from docx.shared import Pt, RGBColor, Cm, Inches
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.enum.table import WD_TABLE_ALIGNMENT, WD_ALIGN_VERTICAL
from docx.oxml.ns import qn
from docx.oxml import OxmlElement
import datetime

doc = Document()

# ── 페이지 여백 설정 ──────────────────────────────────
section = doc.sections[0]
section.page_width  = Cm(21)
section.page_height = Cm(29.7)
section.top_margin    = Cm(2.5)
section.bottom_margin = Cm(2.5)
section.left_margin   = Cm(3.0)
section.right_margin  = Cm(2.5)

# ── 스타일 헬퍼 함수 ─────────────────────────────────
def set_font(run, name="맑은 고딕", size=11, bold=False,
             color=None, italic=False):
    run.font.name = name
    run.font.size = Pt(size)
    run.font.bold = bold
    run.font.italic = italic
    run.font.color.rgb = RGBColor(*color) if color else RGBColor(0x1F, 0x1F, 0x1F)
    r = run._r
    rPr = r.get_or_add_rPr()
    rFonts = OxmlElement("w:rFonts")
    rFonts.set(qn("w:eastAsia"), name)
    rPr.insert(0, rFonts)

def heading1(text, color=(0x1A, 0x56, 0x9B)):
    p = doc.add_paragraph()
    p.paragraph_format.space_before = Pt(18)
    p.paragraph_format.space_after  = Pt(6)
    r = p.add_run(text)
    set_font(r, size=16, bold=True, color=color)
    # 하단 보더
    pPr = p._p.get_or_add_pPr()
    pb  = OxmlElement("w:pBdr")
    bot = OxmlElement("w:bottom")
    bot.set(qn("w:val"), "single"); bot.set(qn("w:sz"), "6")
    bot.set(qn("w:space"), "1");   bot.set(qn("w:color"), "%02X%02X%02X" % color)
    pb.append(bot); pPr.append(pb)
    return p

def heading2(text, color=(0x2E, 0x75, 0xB6)):
    p = doc.add_paragraph()
    p.paragraph_format.space_before = Pt(12)
    p.paragraph_format.space_after  = Pt(4)
    r = p.add_run("▶  " + text)
    set_font(r, size=13, bold=True, color=color)
    return p

def heading3(text):
    p = doc.add_paragraph()
    p.paragraph_format.space_before = Pt(8)
    p.paragraph_format.space_after  = Pt(3)
    r = p.add_run("◆ " + text)
    set_font(r, size=11, bold=True, color=(0x44, 0x72, 0xC4))
    return p

def body(text, indent=False):
    p = doc.add_paragraph()
    p.paragraph_format.space_after = Pt(3)
    if indent:
        p.paragraph_format.left_indent = Cm(0.5)
    r = p.add_run(text)
    set_font(r, size=10.5)
    return p

def bullet(text, level=0):
    p = doc.add_paragraph(style="List Bullet")
    p.paragraph_format.space_after = Pt(2)
    p.paragraph_format.left_indent = Cm(0.5 + level * 0.5)
    r = p.add_run(text)
    set_font(r, size=10.5)
    return p

def add_colored_cell(cell, text, bg_rgb=(0x2E, 0x75, 0xB6),
                     fg_rgb=(0xFF, 0xFF, 0xFF), bold=True, size=10, center=True):
    tc   = cell._tc
    tcPr = tc.get_or_add_tcPr()
    shd  = OxmlElement("w:shd")
    shd.set(qn("w:val"), "clear")
    shd.set(qn("w:color"), "auto")
    shd.set(qn("w:fill"), "%02X%02X%02X" % bg_rgb)
    tcPr.append(shd)
    cell.vertical_alignment = WD_ALIGN_VERTICAL.CENTER
    p = cell.paragraphs[0]
    if center:
        p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    p.clear()
    r = p.add_run(text)
    set_font(r, size=size, bold=bold, color=fg_rgb)

def add_data_cell(cell, text, bg_rgb=None, bold=False, size=10, center=True):
    if bg_rgb:
        tc   = cell._tc
        tcPr = tc.get_or_add_tcPr()
        shd  = OxmlElement("w:shd")
        shd.set(qn("w:val"), "clear")
        shd.set(qn("w:color"), "auto")
        shd.set(qn("w:fill"), "%02X%02X%02X" % bg_rgb)
        tcPr.append(shd)
    cell.vertical_alignment = WD_ALIGN_VERTICAL.CENTER
    p = cell.paragraphs[0]
    if center:
        p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    p.clear()
    r = p.add_run(text)
    set_font(r, size=size, bold=bold,
             color=(0x1F,0x1F,0x1F) if not bg_rgb else (0x1F,0x1F,0x1F))

def set_col_width(table, col_idx, width_cm):
    for row in table.rows:
        row.cells[col_idx].width = Cm(width_cm)

def add_divider(color=(0xBD, 0xD7, 0xEE)):
    p = doc.add_paragraph()
    p.paragraph_format.space_before = Pt(4)
    p.paragraph_format.space_after  = Pt(4)
    pPr = p._p.get_or_add_pPr()
    pb  = OxmlElement("w:pBdr")
    bot = OxmlElement("w:bottom")
    bot.set(qn("w:val"), "single"); bot.set(qn("w:sz"), "4")
    bot.set(qn("w:space"), "1");   bot.set(qn("w:color"), "%02X%02X%02X" % color)
    pb.append(bot); pPr.append(pb)

# ══════════════════════════════════════════════════════════════
# 표지
# ══════════════════════════════════════════════════════════════
p = doc.add_paragraph()
p.alignment = WD_ALIGN_PARAGRAPH.CENTER
p.paragraph_format.space_before = Pt(60)
r = p.add_run("애니메이션 영화 데이터셋")
set_font(r, size=22, bold=True, color=(0x1A,0x56,0x9B))

p = doc.add_paragraph()
p.alignment = WD_ALIGN_PARAGRAPH.CENTER
r = p.add_run("PyTorch 머신러닝 3종 프로젝트")
set_font(r, size=18, bold=True, color=(0x2E,0x75,0xB6))

p = doc.add_paragraph()
p.alignment = WD_ALIGN_PARAGRAPH.CENTER
r = p.add_run("분석 결과 보고서")
set_font(r, size=16, color=(0x44,0x72,0xC4))

doc.add_paragraph()
p = doc.add_paragraph()
p.alignment = WD_ALIGN_PARAGRAPH.CENTER
r = p.add_run("─" * 30)
set_font(r, size=12, color=(0x9D,0xC3,0xE6))

for _ in range(3):
    doc.add_paragraph()

# 메타 정보 테이블
tbl = doc.add_table(rows=4, cols=2)
tbl.alignment = WD_TABLE_ALIGNMENT.CENTER
tbl.style = "Table Grid"
meta = [
    ("데이터셋", "animation_movies_enriched_1878_2029.csv  (25,390개 영화)"),
    ("분석 일자", datetime.date.today().strftime("%Y년 %m월 %d일")),
    ("프레임워크", "PyTorch 2.10  ·  HuggingFace Transformers 5.5  ·  PyG 2.7"),
    ("실행 환경", "CPU (Intel)  ·  Anaconda Python 3.x  ·  Windows 11"),
]
for i, (k, v) in enumerate(meta):
    add_colored_cell(tbl.rows[i].cells[0], k,
                     bg_rgb=(0x2E,0x75,0xB6), size=10)
    add_data_cell(tbl.rows[i].cells[1], v,
                  bg_rgb=(0xF2,0xF7,0xFF) if i%2==0 else (0xE8,0xF2,0xFF),
                  center=False, size=10)
set_col_width(tbl, 0, 3.5); set_col_width(tbl, 1, 11.0)

doc.add_page_break()

# ══════════════════════════════════════════════════════════════
# 1. 데이터셋 개요
# ══════════════════════════════════════════════════════════════
heading1("1. 데이터셋 개요")
body("본 분석에 사용된 데이터셋은 1878년부터 2029년까지의 애니메이션 영화 25,390편에 "
     "대한 정보를 담고 있으며, 총 44개의 컬럼으로 구성되어 있습니다.")

doc.add_paragraph()
tbl = doc.add_table(rows=8, cols=3)
tbl.style = "Table Grid"
tbl.alignment = WD_TABLE_ALIGNMENT.CENTER
headers = ["항목", "내용", "비고"]
for i, h in enumerate(headers):
    add_colored_cell(tbl.rows[0].cells[i], h, bg_rgb=(0x1A,0x56,0x9B))

rows_data = [
    ("총 데이터 수",      "25,390개",             "1878 ~ 2029년"),
    ("컬럼 수",          "44개",                  "수치형·범주형·텍스트·불리언"),
    ("주요 피처",         "장르, 시놉시스, 평점",  "Rating 평균 6.38점"),
    ("시놉시스 보유",     "22,692개 (89.4%)",     "2,698개 결측"),
    ("평점 데이터",       "16,118개 (63.5%)",     "0.5 ~ 10.0점"),
    ("장르 분류",         "상위 8~10개 장르",      "Animation이 가장 빈번"),
    ("Popularity Tier",  "4단계 분류",            "Obscure/Niche/Popular/Blockbuster"),
]
for i, (a, b, c) in enumerate(rows_data):
    bg = (0xF2,0xF7,0xFF) if i%2==0 else (0xFF,0xFF,0xFF)
    add_data_cell(tbl.rows[i+1].cells[0], a, bg_rgb=bg, bold=True, center=False)
    add_data_cell(tbl.rows[i+1].cells[1], b, bg_rgb=bg, center=False)
    add_data_cell(tbl.rows[i+1].cells[2], c, bg_rgb=bg, center=False)
set_col_width(tbl, 0, 3.5); set_col_width(tbl, 1, 5.0); set_col_width(tbl, 2, 6.0)

doc.add_paragraph()
body("세 가지 머신러닝 접근법을 통해 아래와 같은 분석을 수행하였습니다.")
bullet("프로젝트 1 (NLP): 시놉시스 텍스트 → 장르 예측 및 유사 영화 추천")
bullet("프로젝트 2 (정형 데이터): 수치·범주형 피처 → 평점 회귀 및 흥행작 분류")
bullet("프로젝트 3 (그래프): 영화-감독-성우-장르 관계망 → 노드 분류 및 링크 예측")

doc.add_page_break()

# ══════════════════════════════════════════════════════════════
# 2. 프로젝트 1 - NLP
# ══════════════════════════════════════════════════════════════
heading1("2. 프로젝트 1 — NLP: 시놉시스 기반 장르 분류")

heading2("2-1. 모델 구조 및 학습 설정")
tbl = doc.add_table(rows=8, cols=2)
tbl.style = "Table Grid"
tbl.alignment = WD_TABLE_ALIGNMENT.CENTER
add_colored_cell(tbl.rows[0].cells[0], "설정 항목", bg_rgb=(0x2E,0x75,0xB6))
add_colored_cell(tbl.rows[0].cells[1], "내용",      bg_rgb=(0x2E,0x75,0xB6))
cfg = [
    ("사용 모델",   "DistilBERT (distilbert-base-uncased)"),
    ("파라미터 수", "전체 66,561,800  /  학습 38,210,312 (하위 4레이어 동결)"),
    ("입력",       "Overview 시놉시스 텍스트 (Max Length 128 tokens)"),
    ("출력",       "8개 장르 Multi-label 분류 (sigmoid + threshold 0.5)"),
    ("학습 데이터", "Train 6,400 / Val 1,600 (총 8,000개 샘플)"),
    ("학습 설정",  "Epoch 3  ·  Batch 64  ·  LR 3e-5  ·  AdamW  ·  CPU"),
    ("손실 함수",  "BCEWithLogitsLoss"),
]
for i, (k, v) in enumerate(cfg):
    bg = (0xF2,0xF7,0xFF) if i%2==0 else (0xFF,0xFF,0xFF)
    add_data_cell(tbl.rows[i+1].cells[0], k, bg_rgb=bg, bold=True, center=False)
    add_data_cell(tbl.rows[i+1].cells[1], v, bg_rgb=bg, center=False)
set_col_width(tbl, 0, 3.5); set_col_width(tbl, 1, 11.0)

heading2("2-2. 에폭별 학습 경과")
tbl = doc.add_table(rows=4, cols=6)
tbl.style = "Table Grid"
tbl.alignment = WD_TABLE_ALIGNMENT.CENTER
heads = ["에폭", "Train Loss", "F1-micro", "F1-macro", "Hamming Loss", "Jaccard"]
for i, h in enumerate(heads):
    add_colored_cell(tbl.rows[0].cells[i], h, bg_rgb=(0x1A,0x56,0x9B))
epoch_data = [
    ("1 / 3", "0.4689", "66.13%", "12.50%", "0.1280", "65.57%"),
    ("2 / 3", "0.3293", "67.97%", "20.12%", "0.1245", "66.61%"),
    ("3 / 3", "0.2681", "69.41%", "25.53%", "0.1218", "67.43%"),
]
for i, row in enumerate(epoch_data):
    bg = (0xF2,0xF7,0xFF) if i%2==0 else (0xFF,0xFF,0xFF)
    best_bg = (0xE2,0xF0,0xD9)  # 마지막 에폭(최고) 강조
    use_bg  = best_bg if i==2 else bg
    for j, val in enumerate(row):
        add_data_cell(tbl.rows[i+1].cells[j], val, bg_rgb=use_bg,
                      bold=(i==2), center=True)

heading2("2-3. 최종 정확도 측정 결과 (Validation Set)")

# 핵심 지표 하이라이트 박스
tbl = doc.add_table(rows=2, cols=3)
tbl.style = "Table Grid"
tbl.alignment = WD_TABLE_ALIGNMENT.CENTER
kpi = [
    ("F1-micro\n(전체 레이블 균등)", "69.41%", (0x1A,0x56,0x9B)),
    ("Jaccard Similarity\n(예측-실제 겹침)", "67.43%", (0x37,0x86,0x5E)),
    ("Exact Match\n(완전 일치)", "43.69%", (0x84,0x39,0x3C)),
]
for j, (label, val, color) in enumerate(kpi):
    add_colored_cell(tbl.rows[0].cells[j], label, bg_rgb=color, size=9)
    add_colored_cell(tbl.rows[1].cells[j], val,
                     bg_rgb=(0xBD,0xD7,0xEE), fg_rgb=(0x1F,0x1F,0x1F),
                     bold=True, size=16)

doc.add_paragraph()
tbl2 = doc.add_table(rows=7, cols=2)
tbl2.style = "Table Grid"
tbl2.alignment = WD_TABLE_ALIGNMENT.CENTER
add_colored_cell(tbl2.rows[0].cells[0], "지표",  bg_rgb=(0x2E,0x75,0xB6))
add_colored_cell(tbl2.rows[0].cells[1], "값",    bg_rgb=(0x2E,0x75,0xB6))
metrics = [
    ("F1-micro  (전체 레이블 균등 평균)",    "0.6941  (69.41%)"),
    ("F1-macro  (클래스별 단순 평균)",       "0.2553  (25.53%)"),
    ("F1-weighted (샘플 수 가중 평균)",      "0.5738  (57.38%)"),
    ("Hamming Loss (낮을수록 좋음)",         "0.1218"),
    ("Jaccard Similarity (샘플 평균)",      "0.6743  (67.43%)"),
    ("Exact Match  (모든 레이블 완전 일치)", "0.4369  (43.69%)"),
]
for i, (k, v) in enumerate(metrics):
    bg = (0xF2,0xF7,0xFF) if i%2==0 else (0xFF,0xFF,0xFF)
    add_data_cell(tbl2.rows[i+1].cells[0], k, bg_rgb=bg, center=False)
    add_data_cell(tbl2.rows[i+1].cells[1], v, bg_rgb=bg, bold=True, center=True)
set_col_width(tbl2, 0, 8.5); set_col_width(tbl2, 1, 6.0)

heading2("2-4. 장르별 분류 성능")
tbl3 = doc.add_table(rows=10, cols=5)
tbl3.style = "Table Grid"
tbl3.alignment = WD_TABLE_ALIGNMENT.CENTER
for j, h in enumerate(["장르", "Precision", "Recall", "F1-score", "Support"]):
    add_colored_cell(tbl3.rows[0].cells[j], h, bg_rgb=(0x1A,0x56,0x9B))
genre_data = [
    ("Action",          "0.64", "0.17", "0.27", "145"),
    ("Adventure",       "0.55", "0.05", "0.10", "230"),
    ("Animation",       "1.00", "1.00", "1.00", "1,600"),
    ("Comedy",          "0.67", "0.01", "0.01", "384"),
    ("Drama",           "0.00", "0.00", "0.00", "129"),
    ("Family",          "0.68", "0.25", "0.37", "406"),
    ("Fantasy",         "1.00", "0.01", "0.02", "200"),
    ("Science Fiction", "0.62", "0.18", "0.28", "145"),
    ("Micro avg",       "0.95", "0.55", "0.69", "3,239"),
]
for i, row in enumerate(genre_data):
    is_total = (i == 8)
    bg = (0xE2,0xF0,0xD9) if is_total else ((0xF2,0xF7,0xFF) if i%2==0 else (0xFF,0xFF,0xFF))
    for j, val in enumerate(row):
        add_data_cell(tbl3.rows[i+1].cells[j], val, bg_rgb=bg,
                      bold=is_total, center=(j>0))

heading2("2-5. 분석 결과 해석 및 개선 방향")
body("【주요 발견】")
bullet("'Animation' 장르는 데이터의 100%가 해당되어 Precision·Recall 모두 1.00 달성")
bullet("Comedy, Drama, Fantasy 등 소수 레이블은 Recall이 0~1% 수준으로 매우 낮음")
bullet("CPU 환경에서 3에폭 학습으로 소수 장르에 대한 학습이 충분하지 않음")
bullet("F1-micro 기준 69.4%는 Transformer 모델이 시놉시스의 맥락을 의미 있게 파악한 결과")

body("【개선 방향】")
bullet("GPU 환경에서 에폭 6 이상 학습 시 F1-macro 50%↑ 기대")
bullet("pos_weight 조정(클래스 불균형 보정)으로 소수 장르 Recall 개선")
bullet("데이터 증강(역번역, Paraphrasing)으로 소수 장르 샘플 확충")
bullet("RoBERTa 또는 DeBERTa 모델로 교체 시 추가 성능 향상 가능")

doc.add_page_break()

# ══════════════════════════════════════════════════════════════
# 3. 프로젝트 2 - 정형 데이터
# ══════════════════════════════════════════════════════════════
heading1("3. 프로젝트 2 — 정형 데이터: Entity Embedding MLP 평점 예측")

heading2("3-1. 모델 구조 및 학습 설정")
tbl = doc.add_table(rows=8, cols=2)
tbl.style = "Table Grid"
tbl.alignment = WD_TABLE_ALIGNMENT.CENTER
add_colored_cell(tbl.rows[0].cells[0], "설정 항목", bg_rgb=(0x2E,0x75,0xB6))
add_colored_cell(tbl.rows[0].cells[1], "내용",      bg_rgb=(0x2E,0x75,0xB6))
cfg2 = [
    ("모델 구조",   "Entity Embedding → MLP Backbone → Regression/Classification 헤드"),
    ("파라미터 수", "235,832"),
    ("범주형 피처", "Animation_Style, MPAA_Rating, Target_Audience, Era, Popularity_Tier, Original_Language, Primary_Genre (7종)"),
    ("수치형 피처", "Release_Year, Runtime, Log(Vote_Count), Log(Popularity), Log(Budget) + 불리언 4종"),
    ("학습 데이터", "Train 12,894  /  Val 3,224  (총 16,118개 샘플)"),
    ("학습 설정",  "Epoch 60  ·  Batch 256  ·  LR 1e-3  ·  Adam  ·  CosineAnnealingLR"),
    ("손실 함수",  "MSELoss (회귀) + 0.5 × BCEWithLogitsLoss (분류)  →  복합 손실"),
]
for i, (k, v) in enumerate(cfg2):
    bg = (0xF2,0xF7,0xFF) if i%2==0 else (0xFF,0xFF,0xFF)
    add_data_cell(tbl.rows[i+1].cells[0], k, bg_rgb=bg, bold=True, center=False)
    add_data_cell(tbl.rows[i+1].cells[1], v, bg_rgb=bg, center=False)
set_col_width(tbl, 0, 3.5); set_col_width(tbl, 1, 11.0)

heading2("3-2. 에폭별 학습 경과 (10 에폭 간격)")
tbl = doc.add_table(rows=7, cols=5)
tbl.style = "Table Grid"
tbl.alignment = WD_TABLE_ALIGNMENT.CENTER
for j, h in enumerate(["에폭", "MAE", "RMSE", "R²", "AUC"]):
    add_colored_cell(tbl.rows[0].cells[j], h, bg_rgb=(0x1A,0x56,0x9B))
ep_data = [
    ("10",  "1.0786", "1.5445", "6.98%",  "0.6985"),
    ("20",  "1.0608", "1.5202", "9.89%",  "0.7314"),
    ("30",  "1.0437", "1.5099", "11.11%", "0.7467"),
    ("40",  "1.0451", "1.5105", "11.04%", "0.7506"),
    ("50",  "1.0401", "1.5092", "11.18%", "0.7530"),
    ("60 ★","1.0393", "1.5096", "11.14%", "0.7534"),
]
for i, row in enumerate(ep_data):
    best = (i == 5)
    bg = (0xE2,0xF0,0xD9) if best else ((0xF2,0xF7,0xFF) if i%2==0 else (0xFF,0xFF,0xFF))
    for j, val in enumerate(row):
        add_data_cell(tbl.rows[i+1].cells[j], val, bg_rgb=bg,
                      bold=best, center=True)

heading2("3-3. 최종 정확도 측정 결과")

tbl_kpi = doc.add_table(rows=2, cols=4)
tbl_kpi.style = "Table Grid"
tbl_kpi.alignment = WD_TABLE_ALIGNMENT.CENTER
kpi2 = [
    ("MAE (회귀)",     "1.04점",  (0x1A,0x56,0x9B)),
    ("R² (결정계수)",  "11.14%",  (0x37,0x86,0x5E)),
    ("Accuracy (Hit)", "71.56%",  (0x84,0x39,0x3C)),
    ("ROC-AUC",        "75.34%",  (0x7B,0x3F,0x9D)),
]
for j, (label, val, color) in enumerate(kpi2):
    add_colored_cell(tbl_kpi.rows[0].cells[j], label, bg_rgb=color, size=9)
    add_colored_cell(tbl_kpi.rows[1].cells[j], val,
                     bg_rgb=(0xBD,0xD7,0xEE), fg_rgb=(0x1F,0x1F,0x1F),
                     bold=True, size=14)

doc.add_paragraph()
heading3("Task 1: Rating 회귀 (연속값 예측)")
tbl_r = doc.add_table(rows=6, cols=2)
tbl_r.style = "Table Grid"
tbl_r.alignment = WD_TABLE_ALIGNMENT.CENTER
add_colored_cell(tbl_r.rows[0].cells[0], "지표",  bg_rgb=(0x2E,0x75,0xB6))
add_colored_cell(tbl_r.rows[0].cells[1], "값",    bg_rgb=(0x2E,0x75,0xB6))
reg_m = [
    ("MAE (평균 절대 오차)",   "1.0393점"),
    ("RMSE (평균 제곱근 오차)","1.5096점"),
    ("R² (결정계수)",          "0.1114  (11.14%)"),
    ("평균 예측값",            "6.361점  (실제 6.415점)"),
    ("예측 범위",              "3.44 ~ 9.98점"),
]
for i, (k, v) in enumerate(reg_m):
    bg = (0xF2,0xF7,0xFF) if i%2==0 else (0xFF,0xFF,0xFF)
    add_data_cell(tbl_r.rows[i+1].cells[0], k, bg_rgb=bg, center=False)
    add_data_cell(tbl_r.rows[i+1].cells[1], v, bg_rgb=bg, bold=True)
set_col_width(tbl_r, 0, 6.5); set_col_width(tbl_r, 1, 8.0)

doc.add_paragraph()
heading3("Task 2: Hit 분류 (Rating ≥ 7.0)")
tbl_c = doc.add_table(rows=5, cols=2)
tbl_c.style = "Table Grid"
tbl_c.alignment = WD_TABLE_ALIGNMENT.CENTER
add_colored_cell(tbl_c.rows[0].cells[0], "지표",  bg_rgb=(0x2E,0x75,0xB6))
add_colored_cell(tbl_c.rows[0].cells[1], "값",    bg_rgb=(0x2E,0x75,0xB6))
cls_m = [
    ("Accuracy  (정확도)",  "0.7156  (71.56%)"),
    ("ROC-AUC",             "0.7534  (75.34%)"),
    ("Precision (Hit)",     "0.66"),
    ("Recall (Hit)",        "0.37"),
]
for i, (k, v) in enumerate(cls_m):
    bg = (0xF2,0xF7,0xFF) if i%2==0 else (0xFF,0xFF,0xFF)
    add_data_cell(tbl_c.rows[i+1].cells[0], k, bg_rgb=bg, center=False)
    add_data_cell(tbl_c.rows[i+1].cells[1], v, bg_rgb=bg, bold=True)
set_col_width(tbl_c, 0, 6.5); set_col_width(tbl_c, 1, 8.0)

doc.add_paragraph()
heading3("평점 구간별 예측 오차")
tbl_s = doc.add_table(rows=6, cols=4)
tbl_s.style = "Table Grid"
tbl_s.alignment = WD_TABLE_ALIGNMENT.CENTER
for j, h in enumerate(["평점 구간", "샘플 수", "MAE", "RMSE"]):
    add_colored_cell(tbl_s.rows[0].cells[j], h, bg_rgb=(0x1A,0x56,0x9B))
seg_data = [
    ("0 ~ 5점",   "395",   "2.4308", "2.8166"),
    ("5 ~ 6점",   "576",   "0.7168", "0.9006"),
    ("6 ~ 7점 ★", "1,133", "0.4261", "0.5653"),
    ("7 ~ 8점",   "688",   "0.7335", "0.9137"),
    ("8 ~ 10점",  "279",   "1.8194", "2.0179"),
]
for i, row in enumerate(seg_data):
    best = (i == 2)
    bg = (0xE2,0xF0,0xD9) if best else ((0xF2,0xF7,0xFF) if i%2==0 else (0xFF,0xFF,0xFF))
    for j, val in enumerate(row):
        add_data_cell(tbl_s.rows[i+1].cells[j], val, bg_rgb=bg, bold=best)

heading2("3-4. 분석 결과 해석 및 개선 방향")
body("【주요 발견】")
bullet("6~7점 구간(가장 많은 데이터)에서 MAE 0.43으로 가장 정확한 예측")
bullet("극단 구간(0~5점, 8~10점)에서 MAE 1.8~2.4로 오차가 큼 → 데이터 불균형 영향")
bullet("R² 11.14%는 정형 피처(메타데이터)만으로 평점의 변동을 설명하는 데 한계가 있음")
bullet("ROC-AUC 75.34%로 Hit/Not-Hit 이진 분류는 비교적 양호한 변별력을 보임")

body("【개선 방향】")
bullet("Budget, Box_Office 결측치(94% 이상) 보완 시 R² 크게 향상 가능")
bullet("감독 경력, 시리즈 여부 등 추가 파생 피처 생성으로 성능 개선")
bullet("Focal Loss 적용으로 극단 평점 구간의 불균형 문제 해소 가능")
bullet("TabNet, NODE 등 최신 정형 데이터 아키텍처 적용 시 추가 성능 향상 기대")

doc.add_page_break()

# ══════════════════════════════════════════════════════════════
# 4. 프로젝트 3 - 그래프
# ══════════════════════════════════════════════════════════════
heading1("4. 프로젝트 3 — 그래프: Heterogeneous GNN 관계 분석")

heading2("4-1. 그래프 구조 및 학습 설정")
tbl = doc.add_table(rows=9, cols=2)
tbl.style = "Table Grid"
tbl.alignment = WD_TABLE_ALIGNMENT.CENTER
add_colored_cell(tbl.rows[0].cells[0], "설정 항목", bg_rgb=(0x2E,0x75,0xB6))
add_colored_cell(tbl.rows[0].cells[1], "내용",      bg_rgb=(0x2E,0x75,0xB6))
cfg3 = [
    ("그래프 유형", "이분 이종 그래프 (Heterogeneous Bipartite Graph)"),
    ("노드 구성",  "영화 5,000  /  감독 2,576  /  장르 19  /  성우 4,124"),
    ("엣지 구성",  "영화-감독 5,000  /  영화-장르 5,000  /  영화-성우 4,560 (+ 역방향)"),
    ("GNN 모델",  "HeteroConv(GraphSAGE) × 2 레이어  (Hidden 64)"),
    ("Task 1",    "Node Classification — Popularity Tier 4분류"),
    ("Task 2",    "Link Prediction — 감독-성우 협업 가능성 예측"),
    ("학습 설정", "Epoch 100 (NC)  /  80 (LP)  ·  LR 1e-3  ·  Adam  ·  CPU"),
    ("데이터 분할","Train 60%  /  Val 20%  /  Test 20%"),
]
for i, (k, v) in enumerate(cfg3):
    bg = (0xF2,0xF7,0xFF) if i%2==0 else (0xFF,0xFF,0xFF)
    add_data_cell(tbl.rows[i+1].cells[0], k, bg_rgb=bg, bold=True, center=False)
    add_data_cell(tbl.rows[i+1].cells[1], v, bg_rgb=bg, center=False)
set_col_width(tbl, 0, 3.5); set_col_width(tbl, 1, 11.0)

heading2("4-2. 학습 경과")

heading3("Task 1 — Node Classification 학습 경과")
tbl = doc.add_table(rows=6, cols=4)
tbl.style = "Table Grid"
tbl.alignment = WD_TABLE_ALIGNMENT.CENTER
for j, h in enumerate(["에폭", "Loss", "Val Accuracy", "Val F1-weighted"]):
    add_colored_cell(tbl.rows[0].cells[j], h, bg_rgb=(0x1A,0x56,0x9B))
nc_ep = [
    ("20",   "0.8517", "66.90%", "63.57%"),
    ("40",   "0.6438", "71.00%", "70.18%"),
    ("60",   "0.5346", "74.70%", "74.27%"),
    ("80",   "0.4630", "79.60%", "79.55%"),
    ("100★", "0.4008", "81.90%", "81.94%"),
]
for i, row in enumerate(nc_ep):
    best = (i == 4)
    bg = (0xE2,0xF0,0xD9) if best else ((0xF2,0xF7,0xFF) if i%2==0 else (0xFF,0xFF,0xFF))
    for j, val in enumerate(row):
        add_data_cell(tbl.rows[i+1].cells[j], val, bg_rgb=bg, bold=best)

doc.add_paragraph()
heading3("Task 2 — Link Prediction 학습 경과")
tbl = doc.add_table(rows=5, cols=3)
tbl.style = "Table Grid"
tbl.alignment = WD_TABLE_ALIGNMENT.CENTER
for j, h in enumerate(["에폭", "Loss", "Val AUC"]):
    add_colored_cell(tbl.rows[0].cells[j], h, bg_rgb=(0x1A,0x56,0x9B))
lp_ep = [
    ("20",  "0.6734", "0.7512"),
    ("40",  "0.6388", "0.7939"),
    ("60",  "0.5926", "0.8269"),
    ("80★", "0.5420", "0.8512"),
]
for i, row in enumerate(lp_ep):
    best = (i == 3)
    bg = (0xE2,0xF0,0xD9) if best else ((0xF2,0xF7,0xFF) if i%2==0 else (0xFF,0xFF,0xFF))
    for j, val in enumerate(row):
        add_data_cell(tbl.rows[i+1].cells[j], val, bg_rgb=bg, bold=best)

heading2("4-3. 최종 정확도 측정 결과")

tbl_kpi = doc.add_table(rows=2, cols=3)
tbl_kpi.style = "Table Grid"
tbl_kpi.alignment = WD_TABLE_ALIGNMENT.CENTER
kpi3 = [
    ("NC Accuracy\n(노드 분류)",   "78.80%", (0x1A,0x56,0x9B)),
    ("NC F1-macro\n(클래스 평균)", "82.31%", (0x37,0x86,0x5E)),
    ("LP ROC-AUC\n(링크 예측)",    "86.95%", (0x84,0x39,0x3C)),
]
for j, (label, val, color) in enumerate(kpi3):
    add_colored_cell(tbl_kpi.rows[0].cells[j], label, bg_rgb=color, size=9)
    add_colored_cell(tbl_kpi.rows[1].cells[j], val,
                     bg_rgb=(0xBD,0xD7,0xEE), fg_rgb=(0x1F,0x1F,0x1F),
                     bold=True, size=14)

doc.add_paragraph()
heading3("Task 1: Popularity Tier 노드 분류 (Test Set)")
tbl_nc = doc.add_table(rows=6, cols=5)
tbl_nc.style = "Table Grid"
tbl_nc.alignment = WD_TABLE_ALIGNMENT.CENTER
for j, h in enumerate(["클래스", "Precision", "Recall", "F1-score", "Support"]):
    add_colored_cell(tbl_nc.rows[0].cells[j], h, bg_rgb=(0x1A,0x56,0x9B))
nc_data = [
    ("Blockbuster", "0.94", "0.89", "0.91", "18"),
    ("Niche",       "0.75", "0.75", "0.75", "414"),
    ("Obscure",     "0.79", "0.83", "0.81", "446"),
    ("Popular",     "0.89", "0.75", "0.82", "122"),
    ("Weighted avg","0.79", "0.79", "0.79", "1,000"),
]
for i, row in enumerate(nc_data):
    best = (i == 4)
    bg = (0xE2,0xF0,0xD9) if best else ((0xF2,0xF7,0xFF) if i%2==0 else (0xFF,0xFF,0xFF))
    for j, val in enumerate(row):
        add_data_cell(tbl_nc.rows[i+1].cells[j], val, bg_rgb=bg,
                      bold=best, center=(j>0))

doc.add_paragraph()
heading3("Task 1: 혼동 행렬")
tbl_cm = doc.add_table(rows=5, cols=5)
tbl_cm.style = "Table Grid"
tbl_cm.alignment = WD_TABLE_ALIGNMENT.CENTER
add_colored_cell(tbl_cm.rows[0].cells[0], "실제 \\ 예측",   bg_rgb=(0x1A,0x56,0x9B))
for j, cls in enumerate(["Blockbuster","Niche","Obscure","Popular"]):
    add_colored_cell(tbl_cm.rows[0].cells[j+1], cls, bg_rgb=(0x2E,0x75,0xB6), size=9)
cm_data = [
    ("Blockbuster", "16", "0",   "0",   "2"),
    ("Niche",       "0",  "310", "96",  "8"),
    ("Obscure",     "0",  "75",  "370", "1"),
    ("Popular",     "1",  "28",  "1",   "92"),
]
for i, row in enumerate(cm_data):
    add_colored_cell(tbl_cm.rows[i+1].cells[0], row[0],
                     bg_rgb=(0x2E,0x75,0xB6), size=9)
    for j, val in enumerate(row[1:]):
        is_diag = (i == j)
        bg = (0xE2,0xF0,0xD9) if is_diag else (0xFF,0xFF,0xFF)
        add_data_cell(tbl_cm.rows[i+1].cells[j+1], val, bg_rgb=bg, bold=is_diag)

doc.add_paragraph()
heading3("Task 2: 감독-성우 Link Prediction (Test Set)")
tbl_lp = doc.add_table(rows=5, cols=2)
tbl_lp.style = "Table Grid"
tbl_lp.alignment = WD_TABLE_ALIGNMENT.CENTER
add_colored_cell(tbl_lp.rows[0].cells[0], "지표",  bg_rgb=(0x2E,0x75,0xB6))
add_colored_cell(tbl_lp.rows[0].cells[1], "값",    bg_rgb=(0x2E,0x75,0xB6))
lp_m = [
    ("Test Accuracy",  "0.7812  (78.12%)"),
    ("Test ROC-AUC",   "0.8695  (86.95%)"),
    ("Test F1-binary", "0.7904  (79.04%)"),
    ("양성 : 음성 쌍", "4,247 : 4,247  (균형 샘플)"),
]
for i, (k, v) in enumerate(lp_m):
    bg = (0xF2,0xF7,0xFF) if i%2==0 else (0xFF,0xFF,0xFF)
    add_data_cell(tbl_lp.rows[i+1].cells[0], k, bg_rgb=bg, center=False)
    add_data_cell(tbl_lp.rows[i+1].cells[1], v, bg_rgb=bg, bold=True)
set_col_width(tbl_lp, 0, 5.5); set_col_width(tbl_lp, 1, 9.0)

heading2("4-4. 분석 결과 해석 및 개선 방향")
body("【주요 발견】")
bullet("세 프로젝트 중 가장 높은 정확도 달성 (Node Classification 78.8%, Link Prediction AUC 86.9%)")
bullet("Blockbuster 클래스는 샘플이 18개에 불과함에도 Precision 94% — 그래프 관계 특성이 매우 뚜렷")
bullet("Niche-Obscure 간 혼동이 가장 많음 (Niche→Obscure 오분류 96건) — 경계가 모호한 클래스")
bullet("Link Prediction에서 AUC 86.9%는 감독-성우 협업 패턴이 그래프 구조에 잘 반영됨을 의미")

body("【개선 방향】")
bullet("에폭 200+ 및 학습률 웜업 스케줄링 적용 시 노드 분류 85%↑ 기대")
bullet("영화의 수치 피처(평점, 투표수 등)를 감독/성우 노드에도 집계(mean pooling)하여 노드 임베딩 강화")
bullet("HGT(Heterogeneous Graph Transformer)로 교체 시 관계 유형별 어텐션 학습 가능")
bullet("시간축 정보(개봉연도)를 활용한 Temporal GNN으로 트렌드 변화 반영 가능")

doc.add_page_break()

# ══════════════════════════════════════════════════════════════
# 5. 종합 비교
# ══════════════════════════════════════════════════════════════
heading1("5. 프로젝트 종합 비교 및 결론")

heading2("5-1. 핵심 성능 비교표")
tbl = doc.add_table(rows=8, cols=4)
tbl.style = "Table Grid"
tbl.alignment = WD_TABLE_ALIGNMENT.CENTER
for j, h in enumerate(["항목", "프로젝트 1\n(NLP)", "프로젝트 2\n(정형 데이터)", "프로젝트 3\n(그래프)"]):
    add_colored_cell(tbl.rows[0].cells[j], h, bg_rgb=(0x1A,0x56,0x9B))
comp = [
    ("모델",         "DistilBERT",        "Entity Emb. MLP",    "HeteroGNN (SAGE)"),
    ("핵심 지표",    "F1-micro 69.41%",   "Accuracy 71.56%",    "Accuracy 78.80%"),
    ("보조 지표",    "Jaccard 67.43%",    "ROC-AUC 75.34%",     "ROC-AUC 86.95%"),
    ("파라미터",     "66.5M (학습 38M)",  "235K",               "약 500K"),
    ("학습 시간",    "약 21분 (CPU)",     "약 25초 (CPU)",      "약 10초 (CPU)"),
    ("데이터 수",    "8,000개",           "16,118개",           "5,000개 영화"),
    ("강점",         "텍스트 의미 파악",  "학습 속도·안정성",    "관계 구조 학습"),
]
for i, row in enumerate(comp):
    bg = (0xF2,0xF7,0xFF) if i%2==0 else (0xFF,0xFF,0xFF)
    # 최고 성능 셀 강조
    for j, val in enumerate(row):
        is_best = (j == 3 and i in [1,2])  # 그래프가 가장 높음
        cell_bg = (0xE2,0xF0,0xD9) if is_best else bg
        add_data_cell(tbl.rows[i+1].cells[j], val,
                      bg_rgb=cell_bg, bold=(j==0 or is_best), center=(j>0))
set_col_width(tbl, 0, 3.0); set_col_width(tbl, 1, 4.2)
set_col_width(tbl, 2, 4.2); set_col_width(tbl, 3, 4.2)

heading2("5-2. 종합 결론")
body("본 분석은 동일한 애니메이션 영화 데이터셋에 대해 세 가지 이질적인 "
     "PyTorch 기반 머신러닝 접근법을 적용한 결과입니다.")

doc.add_paragraph()
body("① 프로젝트 1 (NLP — DistilBERT)")
bullet("시놉시스 텍스트만으로 장르를 예측하는 복잡한 태스크임에도 F1-micro 69.4% 달성")
bullet("'Animation'처럼 지배적인 장르는 완벽히 예측하나 소수 장르 학습에 더 많은 에폭이 필요")
bullet("GPU 및 전체 데이터셋 활용 시 실용적 수준의 성능 달성 가능")

body("② 프로젝트 2 (정형 데이터 — Entity Embedding MLP)")
bullet("메타데이터 기반 평점 예측은 R² 11%로 한계가 있으나 Hit 분류에서 AUC 75.3% 달성")
bullet("투표수·인기도 등 수치 피처가 평점과 상관관계를 가지나 영화 품질은 측정하기 어려운 잠재 요인에 의해 결정됨")
bullet("빠른 학습 속도(25초)로 실시간 예측 서비스에 적합")

body("③ 프로젝트 3 (그래프 — HeteroGNN)")
bullet("세 접근법 중 가장 높은 정확도 — 영화 간 관계(감독·성우·장르)가 Popularity를 결정하는 핵심 요인임을 확인")
bullet("Link Prediction AUC 86.9%는 특정 감독과 협업할 가능성이 높은 성우를 실제로 추천할 수 있는 수준")
bullet("그래프 구조 데이터가 단순 정형 피처보다 풍부한 정보를 제공함")

heading2("5-3. 향후 연구 방향")
bullet("세 모델의 앙상블: NLP 임베딩 + 정형 피처 + GNN 임베딩을 결합한 통합 예측 모델")
bullet("GPU 환경 구축 후 전체 25,390개 데이터와 더 많은 에폭으로 재학습")
bullet("시계열 분석: 연도별 트렌드 변화를 반영한 Temporal GNN 또는 LSTM 모델")
bullet("다국어 시놉시스 처리: mBERT 또는 XLM-RoBERTa로 비영어권 영화 분류 개선")
bullet("실시간 추천 서비스: 학습된 GNN 임베딩을 벡터 DB에 저장하여 유사 영화 검색 서비스 구현")

# ── 꼬리말 ──────────────────────────────────────────
add_divider()
p = doc.add_paragraph()
p.alignment = WD_ALIGN_PARAGRAPH.CENTER
r = p.add_run(f"생성일: {datetime.datetime.now().strftime('%Y년 %m월 %d일 %H:%M')}  |  "
              f"데이터: animation_movies_enriched_1878_2029.csv  |  "
              f"프레임워크: PyTorch · Transformers · PyG")
set_font(r, size=8, color=(0x99,0x99,0x99))

# ── 저장 ────────────────────────────────────────────
SAVE_PATH = r"C:\Users\USER\OneDrive\Desktop\애니메이션_ML_프로젝트_분석결과.docx"
doc.save(SAVE_PATH)
print(f"[완료] Word 파일 저장: {SAVE_PATH}")
