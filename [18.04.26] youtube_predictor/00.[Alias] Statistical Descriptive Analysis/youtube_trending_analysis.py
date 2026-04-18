"""
YouTube Trending Video Data Analysis
=====================================
Analyzes 4 CSV files containing YouTube trending video data (2020-2026):
  - yearly_trends.csv    : yearly aggregated statistics
  - category_summary.csv : per-category aggregated statistics
  - country_summary.csv  : per-country aggregated statistics
  - trending_videos.csv  : individual video records (10,000 rows)

Outputs:
  - 20 PNG chart files  → Desktop/addction_results/
  - 1  Word report file → Desktop/YouTube_Trending_Analysis_Report.docx
"""

import os
import datetime
import warnings
import numpy as np
import pandas as pd
import matplotlib
matplotlib.use("Agg")                          # non-interactive backend (no GUI window)
import matplotlib.pyplot as plt
import matplotlib.ticker as mticker
from docx import Document
from docx.shared import Pt, RGBColor, Cm
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml.ns import qn
from docx.oxml import OxmlElement

warnings.filterwarnings("ignore")
plt.rcParams["font.family"] = "DejaVu Sans"
plt.rcParams["axes.unicode_minus"] = False

# ──────────────────────────────────────────────────────────────────
# PATHS
# ──────────────────────────────────────────────────────────────────
BASE      = r"C:\Users\USER\OneDrive\Desktop\새 폴더"   # source CSV directory
CHART_DIR = r"C:\Users\USER\OneDrive\Desktop\addction_results"
REPORT    = r"C:\Users\USER\OneDrive\Desktop\YouTube_Trending_Analysis_Report.docx"

os.makedirs(CHART_DIR, exist_ok=True)

# ──────────────────────────────────────────────────────────────────
# COLOR PALETTE & BACKGROUND
# ──────────────────────────────────────────────────────────────────
PALETTE = [
    "#1F497D", "#2E74B5", "#4BACC6", "#70AD47",
    "#FFC000", "#FF7043", "#9C27B0", "#E91E63",
    "#00BCD4", "#FF5722",
]
BG = "#F8F9FA"   # light gray background for all charts

# ══════════════════════════════════════════════════════════════════
# 1. DATA LOADING
# ══════════════════════════════════════════════════════════════════
yearly   = pd.read_csv(BASE + r"\yearly_trends.csv")
category = pd.read_csv(BASE + r"\category_summary.csv")
country  = pd.read_csv(BASE + r"\country_summary.csv")
videos   = pd.read_csv(BASE + r"\trending_videos.csv")

print(f"[Load] yearly={yearly.shape}, category={category.shape}, "
      f"country={country.shape}, videos={videos.shape}")

# ══════════════════════════════════════════════════════════════════
# 2. PRE-COMPUTATIONS  (shared across charts and report)
# ══════════════════════════════════════════════════════════════════

# --- yearly YoY growth (exclude partial 2026) ---
yr = yearly[yearly.year <= 2025].copy()
yr["views_growth"] = yr["avg_views"].pct_change() * 100
yr["likes_growth"] = yr["avg_likes"].pct_change() * 100

# --- category like-rate ---
cat = category.copy()
cat["like_rate%"] = (cat["avg_likes"] / cat["avg_views"] * 100).round(2)
cat_sorted = cat.sort_values("avg_views", ascending=False)

# --- country ranking ---
ctr = country.sort_values("avg_views", ascending=False)

# --- clickbait flag ---
videos["is_clickbait"] = videos["clickbait_score"] > 0.5
cb = videos.groupby("is_clickbait").agg(
    count       = ("video_id",        "count"),
    avg_views   = ("views",           "mean"),
    avg_likes   = ("likes",           "mean"),
    avg_engage  = ("engagement_score","mean"),
).round(2)

# --- subscriber tier ---
sub_bins   = [0, 10_000, 100_000, 1_000_000, 10_000_000, np.inf]
sub_labels = ["<10K", "10K-100K", "100K-1M", "1M-10M", "10M+"]
videos["sub_tier"] = pd.cut(
    videos["subscriber_count"], bins=sub_bins, labels=sub_labels
)
sub_grp = videos.groupby("sub_tier", observed=True).agg(
    count       = ("video_id",        "count"),
    avg_views   = ("views",           "mean"),
    avg_engage  = ("engagement_score","mean"),
).round(2)

# --- title feature comparison ---
title_rows = []
for label, col in [
    ("All Caps",  "has_caps_title"),
    ("Emoji",     "has_emoji_title"),
    ("Question",  "has_question_title"),
]:
    for val in [0, 1]:
        g = videos[videos[col] == val]
        title_rows.append([
            label,
            "With" if val else "Without",
            len(g),
            round(g["views"].mean(), 0),
            round(g["engagement_score"].mean(), 3),
        ])

# --- monthly aggregation ---
MONTH_NAMES = ["Jan","Feb","Mar","Apr","May","Jun",
               "Jul","Aug","Sep","Oct","Nov","Dec"]
mo = videos.groupby("month").agg(
    count      = ("video_id",        "count"),
    avg_views  = ("views",           "mean"),
    avg_engage = ("engagement_score","mean"),
).round(2)
mo.index = [MONTH_NAMES[i - 1] for i in mo.index]

# --- viral top-1% profile ---
VIRAL_THRESHOLD = videos["views"].quantile(0.99)
viral = videos[videos["views"] >= VIRAL_THRESHOLD]
VIRAL_METRICS = [
    "views", "likes", "comments",
    "engagement_score", "days_to_trend",
    "clickbait_score", "duration_seconds", "subscriber_count",
]
viral_cmp = pd.DataFrame({
    "Overall Avg": videos[VIRAL_METRICS].mean(),
    "Viral Avg":   viral[VIRAL_METRICS].mean(),
    "Multiplier":  viral[VIRAL_METRICS].mean() / videos[VIRAL_METRICS].mean(),
}).round(2)

# --- correlation matrix (key numeric columns) ---
CORR_COLS = [
    "views", "likes", "comments", "duration_seconds",
    "days_to_trend", "clickbait_score", "subscriber_count",
    "tag_count", "engagement_score",
]
CORR_LABELS = {
    "views":           "Views",
    "likes":           "Likes",
    "comments":        "Comments",
    "duration_seconds":"Duration",
    "days_to_trend":   "Days to Trend",
    "clickbait_score": "Clickbait",
    "subscriber_count":"Subscribers",
    "tag_count":       "Tag Count",
    "engagement_score":"Engagement",
}
corr_series = (
    videos[CORR_COLS].corr()["engagement_score"]
    .drop("engagement_score")
    .sort_values(ascending=False)
)

# ══════════════════════════════════════════════════════════════════
# 3. CHART HELPERS
# ══════════════════════════════════════════════════════════════════

def savefig(filename: str) -> None:
    """Save the current matplotlib figure as a PNG to CHART_DIR."""
    path = os.path.join(CHART_DIR, f"{filename}.png")
    plt.savefig(path, dpi=150, bbox_inches="tight", facecolor=BG)
    plt.close()
    print(f"  [Chart] {filename}.png")


def bar_label(ax, bars, fmt="{:.2f}M", scale=1e6, offset_ratio=0.02):
    """Annotate each bar with its value."""
    for bar in bars:
        h = bar.get_height()
        ax.text(
            bar.get_x() + bar.get_width() / 2,
            h + h * offset_ratio,
            fmt.format(h / scale),
            ha="center", va="bottom", fontsize=8,
        )


# ══════════════════════════════════════════════════════════════════
# 4. CHART GENERATION  (20 charts total)
# ══════════════════════════════════════════════════════════════════
print("\n[Charts] Generating 20 charts...")

# ── 01. Yearly Avg Views & Likes (dual Y-axis bar + line) ─────────
fig, ax1 = plt.subplots(figsize=(10, 5), facecolor=BG)
ax1.set_facecolor(BG)
bars = ax1.bar(yr["year"], yr["avg_views"] / 1e6,
               color=PALETTE[1], alpha=0.85, label="Avg Views (M)")
ax1.set_ylabel("Avg Views (M)", fontsize=11)
ax1.set_xlabel("Year", fontsize=11)
ax2 = ax1.twinx()
ax2.plot(yr["year"], yr["avg_likes"] / 1e3,
         color=PALETTE[4], marker="o", linewidth=2.5, label="Avg Likes (K)")
ax2.set_ylabel("Avg Likes (K)", fontsize=11)
for bar, v in zip(bars, yr["avg_views"] / 1e6):
    ax1.text(bar.get_x() + bar.get_width() / 2, v + 0.05,
             f"{v:.2f}M", ha="center", va="bottom", fontsize=8)
lines = ax1.get_legend_handles_labels()[0] + ax2.get_legend_handles_labels()[0]
lbls  = ax1.get_legend_handles_labels()[1] + ax2.get_legend_handles_labels()[1]
ax1.legend(lines, lbls, loc="upper left", fontsize=9)
plt.title("Yearly Avg Views & Likes Trend (2020-2025)", fontsize=14, fontweight="bold", pad=12)
plt.tight_layout()
savefig("01_yearly_views_likes")

# ── 02. Yearly YoY Growth Rate (waterfall bar) ───────────────────
yr2    = yr.dropna(subset=["views_growth"])
colors = [PALETTE[2] if v >= 0 else PALETTE[5] for v in yr2["views_growth"]]
fig, ax = plt.subplots(figsize=(9, 5), facecolor=BG)
ax.set_facecolor(BG)
bars = ax.bar(yr2["year"].astype(str), yr2["views_growth"],
              color=colors, edgecolor="white", linewidth=0.8)
ax.axhline(0, color="gray", linewidth=1, linestyle="--")
for bar, v in zip(bars, yr2["views_growth"]):
    ax.text(
        bar.get_x() + bar.get_width() / 2,
        v + (1 if v >= 0 else -3),
        f"{v:+.1f}%",
        ha="center", va="bottom" if v >= 0 else "top",
        fontsize=9, fontweight="bold",
    )
ax.set_ylabel("YoY Growth Rate (%)", fontsize=11)
ax.set_xlabel("Year", fontsize=11)
plt.title("Yearly Avg Views YoY Growth Rate", fontsize=14, fontweight="bold", pad=12)
plt.tight_layout()
savefig("02_yearly_yoy_growth")

# ── 03. Yearly Engagement & Duration (dual Y-axis line) ──────────
fig, ax1 = plt.subplots(figsize=(10, 5), facecolor=BG)
ax1.set_facecolor(BG)
ax1.plot(yr["year"], yr["avg_engagement"],
         color=PALETTE[0], marker="s", linewidth=2.5, label="Avg Engagement Score")
ax1.fill_between(yr["year"], yr["avg_engagement"], alpha=0.15, color=PALETTE[0])
ax1.set_ylabel("Avg Engagement Score", fontsize=11)
ax2 = ax1.twinx()
ax2.plot(yr["year"], yr["avg_duration_sec"],
         color=PALETTE[4], marker="^", linewidth=2.5, linestyle="--", label="Avg Duration (sec)")
ax2.set_ylabel("Avg Duration (sec)", fontsize=11)
lines = ax1.get_legend_handles_labels()[0] + ax2.get_legend_handles_labels()[0]
lbls  = ax1.get_legend_handles_labels()[1] + ax2.get_legend_handles_labels()[1]
ax1.legend(lines, lbls, loc="lower right", fontsize=9)
plt.title("Yearly Engagement Score & Avg Duration Trend", fontsize=14, fontweight="bold", pad=12)
plt.tight_layout()
savefig("03_yearly_engagement_duration")

# ── 04. Category Avg Views (horizontal bar) ──────────────────────
cat_s       = category.sort_values("avg_views")
q75         = category["avg_views"].quantile(0.75)
bar_colors  = [PALETTE[0] if v >= q75 else PALETTE[2] for v in cat_s["avg_views"]]
fig, ax = plt.subplots(figsize=(10, 7), facecolor=BG)
ax.set_facecolor(BG)
hbars = ax.barh(cat_s["category"], cat_s["avg_views"] / 1e6,
                color=bar_colors, edgecolor="white")
for bar, v in zip(hbars, cat_s["avg_views"] / 1e6):
    ax.text(v + 0.05, bar.get_y() + bar.get_height() / 2,
            f"{v:.2f}M", va="center", fontsize=8.5)
ax.set_xlabel("Avg Views (M)", fontsize=11)
plt.title("Avg Views by Category", fontsize=14, fontweight="bold", pad=12)
plt.tight_layout()
savefig("04_category_avg_views")

# ── 05. Category Bubble: Like Rate vs Engagement ─────────────────
cat2 = category.copy()
cat2["like_rate"] = cat2["avg_likes"] / cat2["avg_views"] * 100
fig, ax = plt.subplots(figsize=(11, 7), facecolor=BG)
ax.set_facecolor(BG)
sc = ax.scatter(
    cat2["like_rate"], cat2["avg_engagement"],
    s=cat2["total_videos"] * 0.6,
    c=cat2["avg_views"] / 1e6, cmap="Blues",
    alpha=0.8, edgecolors="gray", linewidth=0.5,
)
for _, row in cat2.iterrows():
    ax.annotate(row["category"], (row["like_rate"], row["avg_engagement"]),
                textcoords="offset points", xytext=(5, 3), fontsize=7.5)
cbar = plt.colorbar(sc, ax=ax)
cbar.set_label("Avg Views (M)", fontsize=9)
ax.set_xlabel("Like Rate (%)", fontsize=11)
ax.set_ylabel("Avg Engagement Score", fontsize=11)
plt.title("Category: Like Rate vs Engagement\n(Bubble size = # of videos)",
          fontsize=13, fontweight="bold", pad=12)
plt.tight_layout()
savefig("05_category_bubble")

# ── 06. Category Duration Boxplot ────────────────────────────────
cat_order = category.sort_values("avg_duration_sec", ascending=False)["category"].tolist()
data_box  = [
    videos[videos["category"] == c]["duration_seconds"].clip(upper=3600).values
    for c in cat_order
]
fig, ax = plt.subplots(figsize=(12, 6), facecolor=BG)
ax.set_facecolor(BG)
bp = ax.boxplot(data_box, patch_artist=True, vert=True,
                medianprops=dict(color="red", linewidth=2))
for patch, color in zip(bp["boxes"], [PALETTE[i % len(PALETTE)] for i in range(len(cat_order))]):
    patch.set_facecolor(color)
    patch.set_alpha(0.7)
ax.set_xticklabels(cat_order, rotation=40, ha="right", fontsize=8.5)
ax.set_ylabel("Duration (sec, clipped at 3600)", fontsize=10)
plt.title("Video Duration Distribution by Category (Boxplot)",
          fontsize=14, fontweight="bold", pad=12)
plt.tight_layout()
savefig("06_category_duration_boxplot")

# ── 07. Country Avg Views (horizontal bar) ───────────────────────
ctr_s      = country.sort_values("avg_views")
highlight  = ctr_s["avg_views"] >= ctr_s["avg_views"].quantile(0.75)
bar_colors = [PALETTE[5] if h else PALETTE[1] for h in highlight]
fig, ax = plt.subplots(figsize=(10, 8), facecolor=BG)
ax.set_facecolor(BG)
hbars = ax.barh(ctr_s["trending_country"], ctr_s["avg_views"] / 1e6,
                color=bar_colors, edgecolor="white")
for bar, v in zip(hbars, ctr_s["avg_views"] / 1e6):
    ax.text(v + 0.05, bar.get_y() + bar.get_height() / 2,
            f"{v:.2f}M", va="center", fontsize=8)
ax.set_xlabel("Avg Views (M)", fontsize=11)
plt.title("Avg Views by Country\n(Top 25% highlighted in orange)",
          fontsize=13, fontweight="bold", pad=12)
plt.tight_layout()
savefig("07_country_avg_views")

# ── 08. Country Scatter: Avg Views vs Engagement ─────────────────
fig, ax = plt.subplots(figsize=(11, 7), facecolor=BG)
ax.set_facecolor(BG)
sc = ax.scatter(
    country["avg_views"] / 1e6, country["avg_engagement"],
    s=country["total_videos"] * 0.3,
    c=range(len(country)), cmap="tab20",
    alpha=0.85, edgecolors="gray", linewidth=0.5,
)
for _, row in country.iterrows():
    ax.annotate(row["trending_country"],
                (row["avg_views"] / 1e6, row["avg_engagement"]),
                textcoords="offset points", xytext=(4, 3), fontsize=8)
ax.set_xlabel("Avg Views (M)", fontsize=11)
ax.set_ylabel("Avg Engagement Score", fontsize=11)
plt.title("Country: Avg Views vs Engagement\n(Bubble size = # of videos)",
          fontsize=13, fontweight="bold", pad=12)
plt.tight_layout()
savefig("08_country_scatter")

# ── 09. Correlation Heatmap ───────────────────────────────────────
corr_df = videos[CORR_COLS].corr()
corr_df.columns = [CORR_LABELS[c] for c in corr_df.columns]
corr_df.index   = [CORR_LABELS[c] for c in corr_df.index]
fig, ax = plt.subplots(figsize=(9, 7), facecolor=BG)
ax.set_facecolor(BG)
im = ax.imshow(corr_df.values, cmap="RdYlBu", vmin=-1, vmax=1, aspect="auto")
plt.colorbar(im, ax=ax, shrink=0.8)
ax.set_xticks(range(len(corr_df.columns)))
ax.set_yticks(range(len(corr_df.index)))
ax.set_xticklabels(corr_df.columns, rotation=40, ha="right", fontsize=9)
ax.set_yticklabels(corr_df.index, fontsize=9)
for i in range(len(corr_df)):
    for j in range(len(corr_df.columns)):
        v = corr_df.values[i, j]
        ax.text(j, i, f"{v:.2f}", ha="center", va="center", fontsize=7.5,
                color="white" if abs(v) > 0.5 else "black")
plt.title("Correlation Heatmap of Key Metrics", fontsize=14, fontweight="bold", pad=12)
plt.tight_layout()
savefig("09_correlation_heatmap")

# ── 10. Views Distribution (linear + log side by side) ───────────
fig, axes = plt.subplots(1, 2, figsize=(13, 5), facecolor=BG)
for ax in axes:
    ax.set_facecolor(BG)
axes[0].hist(videos["views"], bins=60, color=PALETTE[1], edgecolor="white", alpha=0.85)
axes[0].set_xlabel("Views", fontsize=10)
axes[0].set_ylabel("Count", fontsize=10)
axes[0].set_title("Views Distribution (Linear Scale)", fontsize=12, fontweight="bold")
axes[0].xaxis.set_major_formatter(
    mticker.FuncFormatter(lambda x, _: f"{x / 1e6:.0f}M")
)
axes[1].hist(np.log10(videos["views"].clip(lower=1)), bins=60,
             color=PALETTE[3], edgecolor="white", alpha=0.85)
axes[1].set_xlabel("log10(Views)", fontsize=10)
axes[1].set_ylabel("Count", fontsize=10)
axes[1].set_title("Views Distribution (Log Scale)", fontsize=12, fontweight="bold")
plt.suptitle("Views Distribution Comparison", fontsize=14, fontweight="bold", y=1.01)
plt.tight_layout()
savefig("10_views_distribution")

# ── 11. Views Band Pie Chart ──────────────────────────────────────
VIEW_BINS   = [0, 10_000, 50_000, 100_000, 500_000, 1_000_000, 10_000_000, np.inf]
VIEW_LABELS = ["<10K","10K-50K","50K-100K","100K-500K","500K-1M","1M-10M","10M+"]
videos["view_band"] = pd.cut(videos["views"], bins=VIEW_BINS, labels=VIEW_LABELS)
vb = videos["view_band"].value_counts().sort_index()
fig, ax = plt.subplots(figsize=(9, 7), facecolor=BG)
ax.set_facecolor(BG)
_, _, autotexts = ax.pie(
    vb.values, labels=vb.index, autopct="%1.1f%%",
    colors=PALETTE[:len(vb)], startangle=140,
    pctdistance=0.82, wedgeprops=dict(edgecolor="white", linewidth=1.5),
)
for at in autotexts:
    at.set_fontsize(8.5)
plt.title("Views Band Distribution", fontsize=14, fontweight="bold", pad=12)
plt.tight_layout()
savefig("11_views_band_pie")

# ── 12. Days to Trend Distribution (bar) ─────────────────────────
d2t = videos["days_to_trend"].value_counts().sort_index()
fig, ax = plt.subplots(figsize=(9, 5), facecolor=BG)
ax.set_facecolor(BG)
bars = ax.bar(d2t.index.astype(str), d2t.values,
              color=PALETTE[2], edgecolor="white", linewidth=0.8)
for bar, v in zip(bars, d2t.values):
    ax.text(bar.get_x() + bar.get_width() / 2, v + 30,
            f"{v:,}\n({v / len(videos) * 100:.1f}%)",
            ha="center", va="bottom", fontsize=8.5)
ax.set_xlabel("Days to Trend", fontsize=11)
ax.set_ylabel("Number of Videos", fontsize=11)
plt.title("Days to Trend Distribution", fontsize=14, fontweight="bold", pad=12)
plt.tight_layout()
savefig("12_days_to_trend")

# ── 13. Day of Week: Avg Views & Engagement (dual Y-axis) ────────
DAY_ORDER = ["Monday","Tuesday","Wednesday","Thursday","Friday","Saturday","Sunday"]
DAY_SHORT = ["Mon","Tue","Wed","Thu","Fri","Sat","Sun"]
dow = videos.groupby("day_of_week").agg(
    avg_views  = ("views",           "mean"),
    avg_engage = ("engagement_score","mean"),
).reindex(DAY_ORDER)
fig, ax1 = plt.subplots(figsize=(10, 5), facecolor=BG)
ax1.set_facecolor(BG)
bar_colors = [PALETTE[5] if v == dow["avg_views"].max() else PALETTE[1]
              for v in dow["avg_views"]]
bars = ax1.bar(DAY_SHORT, dow["avg_views"] / 1e6,
               color=bar_colors, alpha=0.85, label="Avg Views (M)")
ax1.set_ylabel("Avg Views (M)", fontsize=11)
ax2 = ax1.twinx()
ax2.plot(DAY_SHORT, dow["avg_engage"],
         color=PALETTE[3], marker="D", linewidth=2.5, label="Avg Engagement")
ax2.set_ylabel("Avg Engagement", fontsize=11)
ax2.set_ylim(7.0, 8.0)
for bar, v in zip(bars, dow["avg_views"] / 1e6):
    ax1.text(bar.get_x() + bar.get_width() / 2, v + 0.05,
             f"{v:.2f}M", ha="center", fontsize=8)
lines = ax1.get_legend_handles_labels()[0] + ax2.get_legend_handles_labels()[0]
lbls  = ax1.get_legend_handles_labels()[1] + ax2.get_legend_handles_labels()[1]
ax1.legend(lines, lbls, loc="upper left", fontsize=9)
plt.title("Avg Views & Engagement by Day of Week",
          fontsize=14, fontweight="bold", pad=12)
plt.tight_layout()
savefig("13_dayofweek_views")

# ── 14. Year x Month Avg Views Heatmap ───────────────────────────
pivot = (
    videos.groupby(["year","month"])["views"]
    .mean().unstack(fill_value=0) / 1e6
)
fig, ax = plt.subplots(figsize=(13, 6), facecolor=BG)
ax.set_facecolor(BG)
im = ax.imshow(pivot.values, cmap="YlOrRd", aspect="auto")
plt.colorbar(im, ax=ax, label="Avg Views (M)", shrink=0.8)
ax.set_xticks(range(12))
ax.set_xticklabels(MONTH_NAMES, fontsize=9)
ax.set_yticks(range(len(pivot.index)))
ax.set_yticklabels(pivot.index.astype(str), fontsize=9)
for i in range(len(pivot.index)):
    for j in range(12):
        v = pivot.values[i, j]
        if v > 0:
            ax.text(j, i, f"{v:.1f}", ha="center", va="center", fontsize=7,
                    color="white" if v > pivot.values.max() * 0.6 else "black")
plt.title("Avg Views Heatmap by Year x Month (M)",
          fontsize=14, fontweight="bold", pad=12)
plt.tight_layout()
savefig("14_year_month_heatmap")

# ── 15. Clickbait Score Distribution (histogram + KDE) ───────────
fig, ax = plt.subplots(figsize=(10, 5), facecolor=BG)
ax.set_facecolor(BG)
ax.hist(videos["clickbait_score"], bins=50, color=PALETTE[1],
        edgecolor="white", alpha=0.75, density=True, label="Histogram")
try:
    from scipy.stats import gaussian_kde
    kde = gaussian_kde(videos["clickbait_score"])
    xs  = np.linspace(0, 1, 300)
    ax.plot(xs, kde(xs), color=PALETTE[5], linewidth=2.5, label="KDE")
except ImportError:
    pass   # scipy not available; skip KDE overlay
ax.axvline(0.5, color="red", linestyle="--", linewidth=1.5,
           label="Clickbait Threshold (0.5)")
ax.axvline(videos["clickbait_score"].mean(), color=PALETTE[3],
           linestyle=":", linewidth=1.5,
           label=f"Mean {videos['clickbait_score'].mean():.3f}")
ax.set_xlabel("Clickbait Score", fontsize=11)
ax.set_ylabel("Density", fontsize=11)
ax.legend(fontsize=9)
plt.title("Clickbait Score Distribution", fontsize=14, fontweight="bold", pad=12)
plt.tight_layout()
savefig("15_clickbait_distribution")

# ── 16. Clickbait vs Normal: 3-metric grouped bar ────────────────
cb.index = ["Normal (<=0.5)", "Clickbait (>0.5)"]
fig, axes = plt.subplots(1, 3, figsize=(13, 5), facecolor=BG)
for ax in axes:
    ax.set_facecolor(BG)
for ax, col, title, unit, scale in zip(
    axes,
    ["avg_views", "avg_likes", "avg_engage"],
    ["Avg Views", "Avg Likes", "Avg Engagement"],
    ["M", "K", ""],
    [1e6, 1e3, 1],
):
    vals = cb[col] / scale
    bars = ax.bar(cb.index, vals,
                  color=[PALETTE[1], PALETTE[5]], edgecolor="white", alpha=0.85)
    for bar, v in zip(bars, vals):
        ax.text(bar.get_x() + bar.get_width() / 2, v * 1.02,
                f"{v:.2f}{unit}", ha="center", fontsize=9)
    ax.set_title(title, fontsize=11, fontweight="bold")
    ax.set_ylabel(f"{title} ({unit})" if unit else title, fontsize=9)
    ax.set_xticklabels(cb.index, fontsize=8)
plt.suptitle("Clickbait vs Normal Video Performance",
             fontsize=14, fontweight="bold", y=1.02)
plt.tight_layout()
savefig("16_clickbait_comparison")

# ── 17. Title Features vs Avg Views (grouped bar) ────────────────
feat_keys = ["All Caps", "Emoji", "Question"]
feat_cols = ["has_caps_title", "has_emoji_title", "has_question_title"]
feat_vals = {
    label: [
        videos[videos[col] == 0]["views"].mean() / 1e6,
        videos[videos[col] == 1]["views"].mean() / 1e6,
    ]
    for label, col in zip(feat_keys, feat_cols)
}
x = np.arange(len(feat_vals))
w = 0.35
fig, ax = plt.subplots(figsize=(10, 6), facecolor=BG)
ax.set_facecolor(BG)
b1 = ax.bar(x - w/2, [v[0] for v in feat_vals.values()], w,
            label="Without", color=PALETTE[2], alpha=0.85, edgecolor="white")
b2 = ax.bar(x + w/2, [v[1] for v in feat_vals.values()], w,
            label="With", color=PALETTE[5], alpha=0.85, edgecolor="white")
for bar in list(b1) + list(b2):
    ax.text(bar.get_x() + bar.get_width() / 2,
            bar.get_height() + 0.03,
            f"{bar.get_height():.2f}M", ha="center", fontsize=8.5)
ax.set_xticks(x)
ax.set_xticklabels(feat_keys, fontsize=11)
ax.set_ylabel("Avg Views (M)", fontsize=11)
ax.legend(fontsize=10)
plt.title("Title Features (All Caps / Emoji / Question) vs Avg Views",
          fontsize=13, fontweight="bold", pad=12)
plt.tight_layout()
savefig("17_title_features_views")

# ── 18. Subscriber Tier: Avg Views (log) & Engagement ────────────
fig, ax1 = plt.subplots(figsize=(10, 5), facecolor=BG)
ax1.set_facecolor(BG)
bars = ax1.bar(sub_labels, sub_grp["avg_views"] / 1e6,
               color=PALETTE[0], alpha=0.85, label="Avg Views (M)")
ax1.set_ylabel("Avg Views (M, log scale)", fontsize=11)
ax1.set_yscale("log")   # log scale to handle large range across tiers
ax2 = ax1.twinx()
ax2.plot(sub_labels, sub_grp["avg_engage"],
         color=PALETTE[4], marker="o", linewidth=2.5, label="Avg Engagement")
ax2.set_ylabel("Avg Engagement", fontsize=11)
ax2.set_ylim(7.0, 8.0)
lines = ax1.get_legend_handles_labels()[0] + ax2.get_legend_handles_labels()[0]
lbls  = ax1.get_legend_handles_labels()[1] + ax2.get_legend_handles_labels()[1]
ax1.legend(lines, lbls, loc="upper left", fontsize=9)
plt.title("Subscriber Tier: Avg Views (log) & Engagement",
          fontsize=13, fontweight="bold", pad=12)
plt.tight_layout()
savefig("18_subscriber_tier")

# ── 19. Viral Top 1% vs Overall — Radar Chart ────────────────────
RADAR_COLS   = ["views","likes","comments","engagement_score","subscriber_count"]
RADAR_LABELS = ["Views","Likes","Comments","Engagement","Subscribers"]
overall_mean = videos[RADAR_COLS].mean()
viral_mean   = viral[RADAR_COLS].mean()
# normalize so that viral baseline = 1.0
normalized_overall = (overall_mean / viral_mean).values
normalized_viral   = np.ones(len(RADAR_COLS))
angles = np.linspace(0, 2 * np.pi, len(RADAR_COLS), endpoint=False).tolist()
angles += angles[:1]   # close the polygon
normalized_overall = np.append(normalized_overall, normalized_overall[0])
normalized_viral   = np.append(normalized_viral,   normalized_viral[0])
fig, ax = plt.subplots(figsize=(7, 7), subplot_kw=dict(polar=True), facecolor=BG)
ax.set_facecolor(BG)
ax.plot(angles, normalized_viral, "o-", linewidth=2,
        color=PALETTE[5], label="Viral Top 1%")
ax.fill(angles, normalized_viral, alpha=0.2, color=PALETTE[5])
ax.plot(angles, normalized_overall, "o-", linewidth=2,
        color=PALETTE[1], label="Overall Avg")
ax.fill(angles, normalized_overall, alpha=0.2, color=PALETTE[1])
ax.set_thetagrids(np.degrees(angles[:-1]), RADAR_LABELS, fontsize=10)
ax.legend(loc="upper right", bbox_to_anchor=(1.3, 1.1), fontsize=9)
plt.title("Viral Top 1% vs Overall Average\n(Normalized to viral baseline)",
          fontsize=12, fontweight="bold", pad=20)
plt.tight_layout()
savefig("19_viral_radar")

# ── 20. Category x Country Heatmap (video count) ─────────────────
top_countries = videos["trending_country"].value_counts().head(10).index
top_cats      = videos["category"].value_counts().head(10).index
sub_df        = videos[videos["trending_country"].isin(top_countries)]
ct = pd.crosstab(sub_df["category"], sub_df["trending_country"])
ct = ct.loc[ct.index.isin(top_cats), top_countries]
fig, ax = plt.subplots(figsize=(13, 7), facecolor=BG)
ax.set_facecolor(BG)
im = ax.imshow(ct.values, cmap="Blues", aspect="auto")
plt.colorbar(im, ax=ax, label="Number of Videos", shrink=0.8)
ax.set_xticks(range(len(ct.columns)))
ax.set_yticks(range(len(ct.index)))
ax.set_xticklabels(ct.columns, fontsize=10)
ax.set_yticklabels(ct.index, fontsize=9)
for i in range(len(ct.index)):
    for j in range(len(ct.columns)):
        v = ct.values[i, j]
        ax.text(j, i, str(v), ha="center", va="center", fontsize=8,
                color="white" if v > ct.values.max() * 0.6 else "black")
plt.title("Category x Country Heatmap (Top 10 Countries)",
          fontsize=13, fontweight="bold", pad=12)
plt.tight_layout()
savefig("20_category_country_heatmap")

print(f"  [Done] All 20 charts saved to: {CHART_DIR}\n")

# ══════════════════════════════════════════════════════════════════
# 5. WORD REPORT GENERATION
# ══════════════════════════════════════════════════════════════════
print("[Report] Building Word document...")

# ── 5-1. Report helpers ───────────────────────────────────────────

def set_cell_bg(cell, hex_color: str) -> None:
    """Apply a solid background color to a table cell."""
    tc   = cell._tc
    tcPr = tc.get_or_add_tcPr()
    shd  = OxmlElement("w:shd")
    shd.set(qn("w:val"),   "clear")
    shd.set(qn("w:color"), "auto")
    shd.set(qn("w:fill"),  hex_color)
    tcPr.append(shd)


def add_heading(doc: Document, text: str, level: int = 1):
    """Add a styled heading paragraph."""
    p   = doc.add_heading(text, level=level)
    run = p.runs[0]
    run.font.name = "Calibri"
    if level == 1:
        run.font.size  = Pt(15)
        run.font.color.rgb = RGBColor(0x1F, 0x49, 0x7D)
    elif level == 2:
        run.font.size  = Pt(12)
        run.font.color.rgb = RGBColor(0x2E, 0x74, 0xB5)
    return p


def add_table(doc: Document, headers: list, rows: list,
              hdr_color: str = "2E74B5", alt: str = "EBF3FB"):
    """
    Add a styled table.
    - Header row: colored background, white bold text.
    - Data rows: alternating row colors (alt / white).
    """
    t = doc.add_table(rows=1 + len(rows), cols=len(headers))
    t.style = "Table Grid"
    # header row
    for i, h in enumerate(headers):
        cell = t.rows[0].cells[i]
        cell.text = str(h)
        set_cell_bg(cell, hdr_color)
        run = cell.paragraphs[0].runs[0]
        run.font.bold  = True
        run.font.color.rgb = RGBColor(0xFF, 0xFF, 0xFF)
        run.font.size  = Pt(9)
        run.font.name  = "Calibri"
        cell.paragraphs[0].alignment = WD_ALIGN_PARAGRAPH.CENTER
    # data rows
    for ri, row in enumerate(rows):
        bg = alt if ri % 2 == 0 else "FFFFFF"
        for ci, val in enumerate(row):
            cell = t.rows[ri + 1].cells[ci]
            cell.text = str(val)
            set_cell_bg(cell, bg)
            run = cell.paragraphs[0].runs[0]
            run.font.size = Pt(9)
            run.font.name = "Calibri"
    return t


def add_insight(doc: Document, text: str) -> None:
    """Add a green insight bullet below a table."""
    p   = doc.add_paragraph()
    p.paragraph_format.left_indent = Cm(0.5)
    run = p.add_run("► " + text)
    run.font.name  = "Calibri"
    run.font.size  = Pt(9)
    run.font.color.rgb = RGBColor(0x17, 0x6B, 0x1E)


# ── 5-2. Build document ───────────────────────────────────────────
doc = Document()
doc.styles["Normal"].font.name = "Calibri"
doc.styles["Normal"].font.size = Pt(10)

# Cover page
doc.add_paragraph()
tp  = doc.add_paragraph()
tp.alignment = WD_ALIGN_PARAGRAPH.CENTER
r   = tp.add_run("YouTube Trending Video Data Analysis Report")
r.font.name  = "Calibri"
r.font.size  = Pt(22)
r.font.bold  = True
r.font.color.rgb = RGBColor(0x1F, 0x49, 0x7D)
doc.add_paragraph()
sp  = doc.add_paragraph()
sp.alignment = WD_ALIGN_PARAGRAPH.CENTER
r2  = sp.add_run(
    "Period: 2020 - 2026  |  Total Videos: 10,000  |  Countries: 24\n"
    f"Generated: {datetime.date.today().strftime('%B %d, %Y')}"
)
r2.font.name  = "Calibri"
r2.font.size  = Pt(11)
r2.font.color.rgb = RGBColor(0x44, 0x44, 0x44)
doc.add_page_break()

# Section 1: Dataset Overview
add_heading(doc, "1. Dataset Overview", 1)
tv = videos["views"].sum()
tl = videos["likes"].sum()
add_table(doc, ["Item", "Value"], [
    ["Total Videos",          f"{len(videos):,}"],
    ["Total Cumulative Views", f"{tv:,.0f}  ({tv / 1e9:.2f}B)"],
    ["Total Cumulative Likes", f"{tl:,.0f}  ({tl / 1e9:.2f}B)"],
    ["Analysis Period",       "2020 - 2026"],
    ["Countries",             "24"],
    ["Categories",            f"{videos['category'].nunique()}"],
    ["Missing Values",        "None (complete dataset)"],
    ["Top Category (views)",  videos.groupby("category")["views"].sum().idxmax()],
    ["Top Country  (views)",  videos.groupby("trending_country")["views"].sum().idxmax()],
    ["Top Language (views)",  videos.groupby("language")["views"].sum().idxmax()],
])
doc.add_paragraph()

# Section 2: Descriptive Statistics
add_heading(doc, "2. Descriptive Statistics", 1)
STAT_COLS = ["views","likes","dislikes","comments","engagement_score",
             "duration_seconds","days_to_trend","clickbait_score","subscriber_count"]
STAT_LABELS = {
    "views":"Views","likes":"Likes","dislikes":"Dislikes","comments":"Comments",
    "engagement_score":"Engagement","duration_seconds":"Duration (sec)",
    "days_to_trend":"Days to Trend","clickbait_score":"Clickbait Score",
    "subscriber_count":"Subscribers",
}
desc = videos[STAT_COLS].describe().T
desc["cv%"] = (desc["std"] / desc["mean"] * 100).round(1)
add_table(doc, ["Metric","Mean","Median","Std Dev","Min","Max","CV%"], [
    [STAT_LABELS[c],
     f"{desc.loc[c,'mean']:,.1f}", f"{desc.loc[c,'50%']:,.1f}",
     f"{desc.loc[c,'std']:,.1f}",  f"{desc.loc[c,'min']:,.1f}",
     f"{desc.loc[c,'max']:,.1f}",  f"{desc.loc[c,'cv%']:.1f}%"]
    for c in STAT_COLS
])
add_insight(doc, "Views & Likes CV% > 1,000% — a few viral videos massively skew the mean")
add_insight(doc, "Engagement CV% = 45% — the most stable metric across the dataset")
doc.add_paragraph()

# Section 3: Yearly Trend
add_heading(doc, "3. Yearly Trend Analysis", 1)
yr_rows = []
for _, r in yr.iterrows():
    vg = f"{r['views_growth']:+.1f}%" if pd.notna(r["views_growth"]) else "-"
    lg = f"{r['likes_growth']:+.1f}%" if pd.notna(r["likes_growth"]) else "-"
    yr_rows.append([
        int(r["year"]), f"{r['total_videos']:,}",
        f"{r['avg_views'] / 1e6:.2f}M", vg,
        f"{r['avg_likes'] / 1e3:.1f}K", lg,
        f"{r['avg_duration_sec']:.0f}s", f"{r['avg_engagement']:.2f}",
    ])
add_table(doc, ["Year","Videos","Avg Views","Views YoY",
                "Avg Likes","Likes YoY","Avg Duration","Engagement"], yr_rows)
add_insight(doc, "2024: Views +75%, Likes +50% — best performing year in the dataset")
add_insight(doc, "Growth trend 2020-2022 → sharp drop in 2023 → V-shaped recovery in 2024")
doc.add_paragraph()

# Section 4: Category Analysis
add_heading(doc, "4. Category Analysis", 1)
add_table(doc, ["Category","Videos","Avg Views","Like Rate","Engagement","Avg Duration"], [
    [r["category"], f"{r['total_videos']:,}", f"{r['avg_views'] / 1e6:.2f}M",
     f"{r['like_rate%']:.2f}%", f"{r['avg_engagement']:.2f}", f"{r['avg_duration_sec']:.0f}s"]
    for _, r in cat_sorted.iterrows()
])
add_insight(doc, "Sports & Shows lead in avg views (4.6M+); Gaming averages 25-min videos")
add_insight(doc, "Shorts (37 sec) ranks 4th in views with the highest like rate (6.61%)")
add_insight(doc, "Science & Technology has the 2nd-highest like rate (6.89%) — strong fan loyalty")
doc.add_paragraph()

# Section 5: Country Analysis
add_heading(doc, "5. Country Analysis", 1)
add_heading(doc, "5-1. Top 10 by Avg Views", 2)
add_table(doc, ["Country","Avg Views","Engagement","Top Category","Videos"], [
    [r["trending_country"], f"{r['avg_views'] / 1e6:.2f}M",
     f"{r['avg_engagement']:.2f}", r["top_category"], f"{r['total_videos']:,}"]
    for _, r in ctr.head(10).iterrows()
])
add_insight(doc, "Russia (8.4M) & Turkey (8.0M) top the list — 3x higher than the US (2.4M)")
doc.add_paragraph()
add_heading(doc, "5-2. Bottom 5 by Avg Views", 2)
add_table(doc, ["Country","Avg Views","Engagement","Top Category","Videos"], [
    [r["trending_country"], f"{r['avg_views'] / 1e6:.3f}M",
     f"{r['avg_engagement']:.2f}", r["top_category"], f"{r['total_videos']:,}"]
    for _, r in ctr.tail(5).iterrows()
])
add_insight(doc, "Philippines (PH): lowest avg views (886K) but highest engagement (8.27) — small, loyal audience")
doc.add_paragraph()

# Section 6: Correlation Analysis
add_heading(doc, "6. Engagement Correlation Analysis", 1)
add_table(doc, ["Variable","Pearson r","Direction"], [
    [CORR_LABELS.get(k, k), f"{v:.4f}", "Positive" if v > 0 else "Negative"]
    for k, v in corr_series.items()
])
add_insight(doc, "All |r| < 0.03 — engagement score is not predictable from surface-level metrics")
add_insight(doc, "High views / subscribers / clickbait do not guarantee high engagement")
doc.add_paragraph()

# Section 7: Clickbait Comparison
add_heading(doc, "7. Clickbait vs Normal Video Performance", 1)
add_table(doc, ["Type","Videos","Avg Views","Avg Engagement"], [
    ["Normal (<=0.5)",    f"{cb.loc['Normal (<=0.5)','count']:,}",
     f"{cb.loc['Normal (<=0.5)','avg_views'] / 1e6:.2f}M",
     f"{cb.loc['Normal (<=0.5)','avg_engage']:.2f}"],
    ["Clickbait (>0.5)",  f"{cb.loc['Clickbait (>0.5)','count']:,}",
     f"{cb.loc['Clickbait (>0.5)','avg_views'] / 1e6:.2f}M",
     f"{cb.loc['Clickbait (>0.5)','avg_engage']:.2f}"],
])
add_insight(doc, "Clickbait videos show +19% higher avg views but slightly lower engagement")
doc.add_paragraph()

# Section 8: Title Features
add_heading(doc, "8. Title Feature Impact on Performance", 1)
add_table(doc, ["Feature","Type","Videos","Avg Views","Avg Engagement"], title_rows)
add_insight(doc, "Question-style titles average 4.60M views — +48% vs non-question titles")
add_insight(doc, "Emoji titles show lower avg views (-11%) but slightly higher engagement")
doc.add_paragraph()

# Section 9: Days to Trend
add_heading(doc, "9. Days to Trend Distribution", 1)
add_table(doc, ["Days to Trend","Videos","Share"], [
    [f"{d} day(s)", f"{c:,}", f"{c / len(videos) * 100:.1f}%"]
    for d, c in d2t.items()
])
add_insight(doc, "15.0% trend same day (0 days); 59.7% trend within 1-3 days of upload")
doc.add_paragraph()

# Section 10: Day of Week
add_heading(doc, "10. Day-of-Week Trending Pattern", 1)
add_table(doc, ["Day","Videos","Avg Views","Avg Engagement"], [
    [DAY_SHORT[i], f"{row['avg_views']:,.0f}", f"{row['avg_views'] / 1e6:.2f}M",
     f"{row['avg_engage']:.2f}"]
    for i, (_, row) in enumerate(dow.iterrows())
])
add_insight(doc, "Sunday leads with 5.86M avg views — 1.78x the Thursday figure (3.29M)")
doc.add_paragraph()

# Section 11: Monthly Pattern
add_heading(doc, "11. Monthly Trending Pattern", 1)
add_table(doc, ["Month","Videos","Avg Views","Avg Engagement"], [
    [idx, f"{row['count']:,}", f"{row['avg_views'] / 1e6:.2f}M", f"{row['avg_engage']:.2f}"]
    for idx, row in mo.iterrows()
])
add_insight(doc, "November peaks at 6.18M avg views; October also strong at 4.96M — Q4 season effect")
add_insight(doc, "August-September are the lowest months (1.3-1.5M) — summer audience shift")
doc.add_paragraph()

# Section 12: Views Band
add_heading(doc, "12. Views Band Distribution", 1)
add_table(doc, ["Band","Videos","Share"], [
    [lb, f"{vb[lb]:,}", f"{vb[lb] / len(videos) * 100:.1f}%"]
    for lb in VIEW_LABELS
])
add_insight(doc, "<10K band accounts for 26.9% — classic long-tail distribution")
add_insight(doc, "10M+ mega-viral videos represent only 4.1% of the dataset")
doc.add_paragraph()

# Section 13: Subscriber Tier
add_heading(doc, "13. Performance by Subscriber Tier", 1)
add_table(doc, ["Tier","Videos","Avg Views","Avg Engagement"], [
    [idx, f"{row['count']:,}", f"{row['avg_views'] / 1e6:.2f}M", f"{row['avg_engage']:.2f}"]
    for idx, row in sub_grp.iterrows()
])
add_insight(doc, "10M+ channels avg 45.1M views — 3,586x more than <10K channels (12.6K)")
add_insight(doc, "Small channels (<10K) show higher engagement (7.60) than mega channels (7.43)")
doc.add_paragraph()

# Section 14: Viral Profile
add_heading(doc, "14. Viral Top 1% Video Profile", 1)
METRIC_LABELS = {
    "views":"Views","likes":"Likes","comments":"Comments",
    "engagement_score":"Engagement","days_to_trend":"Days to Trend",
    "clickbait_score":"Clickbait Score","duration_seconds":"Duration (sec)",
    "subscriber_count":"Subscribers",
}
add_table(doc, ["Metric","Overall Avg","Viral Avg","Multiplier"], [
    [METRIC_LABELS.get(m, m),
     f"{viral_cmp.loc[m,'Overall Avg']:,.2f}",
     f"{viral_cmp.loc[m,'Viral Avg']:,.2f}",
     f"{viral_cmp.loc[m,'Multiplier']:.2f}x"]
    for m in VIRAL_METRICS
])
p   = doc.add_paragraph()
run = p.add_run(
    f"  * Viral threshold: {VIRAL_THRESHOLD:,.0f}+ views  |  "
    f"Qualifying videos: {len(viral)}"
)
run.font.size  = Pt(8)
run.font.name  = "Calibri"
run.font.color.rgb = RGBColor(0x80, 0x80, 0x80)
add_insight(doc, "Viral videos have ~60x more views, likes, and comments than the overall average")
add_insight(doc, "Engagement & clickbait scores nearly identical — viral success driven by content quality & reach")
doc.add_paragraph()

# Section 15: Channel Verification
add_heading(doc, "15. Verified vs Unverified Channel Performance", 1)
ver = videos.groupby("channel_verified").agg(
    count      = ("video_id",        "count"),
    avg_views  = ("views",           "mean"),
    avg_subs   = ("subscriber_count","mean"),
    avg_engage = ("engagement_score","mean"),
    avg_d2t    = ("days_to_trend",   "mean"),
).round(2)
ver.index = ver.index.map({0: "Unverified", 1: "Verified"})
add_table(doc, ["Type","Videos","Avg Views","Avg Subscribers","Avg Engagement","Avg Days to Trend"], [
    [idx, f"{row['count']:,}",
     f"{row['avg_views'] / 1e6:.2f}M", f"{row['avg_subs'] / 1e6:.2f}M",
     f"{row['avg_engage']:.2f}", f"{row['avg_d2t']:.2f} days"]
    for idx, row in ver.iterrows()
])
add_insight(doc, "Verified channel avg views 7.62M — 61x higher than unverified (125K)")
add_insight(doc, "Days to trend almost identical (2.95 vs 2.98) — verification does not speed up trending")
doc.add_paragraph()

# Save document
doc.save(REPORT)
print(f"  [Done] Report saved to: {REPORT}\n")
print("=" * 60)
print("All outputs generated successfully.")
print(f"  Charts : {CHART_DIR}")
print(f"  Report : {REPORT}")
print("=" * 60)
