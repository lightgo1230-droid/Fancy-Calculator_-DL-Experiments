"""
build_comparison_report.py
Comparative analysis: NumPy/Pandas Statistical vs Deep Learning Approach
English-only Word report
"""

import datetime
from docx import Document
from docx.shared import Inches, Pt, RGBColor, Cm
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.enum.table import WD_TABLE_ALIGNMENT, WD_ALIGN_VERTICAL
from docx.oxml.ns import qn
from docx.oxml import OxmlElement

OUT_PATH = r"C:\Users\USER\OneDrive\Desktop\결과물\Comparative_Analysis_Report.docx"

# ── Colours ────────────────────────────────────────────────────
WHITE   = RGBColor(0xFF, 0xFF, 0xFF)
BLACK   = RGBColor(0x1F, 0x2D, 0x3D)
BLUE    = RGBColor(0x1E, 0x40, 0xAF)
LBLUE   = RGBColor(0x1D, 0x4E, 0xD8)
CYAN    = RGBColor(0x06, 0x7B, 0xA9)
GREEN   = RGBColor(0x06, 0x5F, 0x46)
LGREEN  = RGBColor(0x05, 0x96, 0x69)
RED     = RGBColor(0x99, 0x1B, 0x1B)
SLATE   = RGBColor(0x47, 0x55, 0x69)
MUTED   = RGBColor(0x6B, 0x72, 0x80)

# ── XML helpers ────────────────────────────────────────────────
def set_cell_bg(cell, hex_color):
    tc   = cell._tc
    tcPr = tc.get_or_add_tcPr()
    shd  = OxmlElement('w:shd')
    shd.set(qn('w:val'),   'clear')
    shd.set(qn('w:color'), 'auto')
    shd.set(qn('w:fill'),  hex_color)
    tcPr.append(shd)

def set_row_h(row, cm):
    tr   = row._tr
    trPr = tr.get_or_add_trPr()
    trH  = OxmlElement('w:trHeight')
    trH.set(qn('w:val'), str(int(cm * 567)))
    trPr.append(trH)

def hr_line(doc, color="1E40AF"):
    p    = doc.add_paragraph()
    pPr  = p._p.get_or_add_pPr()
    pBdr = OxmlElement('w:pBdr')
    bot  = OxmlElement('w:bottom')
    bot.set(qn('w:val'), 'single'); bot.set(qn('w:sz'), '6')
    bot.set(qn('w:space'), '1');    bot.set(qn('w:color'), color)
    pBdr.append(bot); pPr.append(pBdr)
    p.paragraph_format.space_before = Pt(2)
    p.paragraph_format.space_after  = Pt(6)

def heading(doc, text, size=14, color=BLUE, bold=True, before=14, after=6):
    p   = doc.add_paragraph()
    run = p.add_run(text)
    run.bold = bold; run.font.size = Pt(size); run.font.color.rgb = color
    p.paragraph_format.space_before = Pt(before)
    p.paragraph_format.space_after  = Pt(after)
    return p

def body(doc, text, size=10.5, color=BLACK, bold=False, after=6, italic=False):
    p   = doc.add_paragraph()
    run = p.add_run(text)
    run.font.size = Pt(size); run.font.color.rgb = color
    run.bold = bold; run.italic = italic
    p.paragraph_format.space_after = Pt(after)
    return p

def bullet(doc, text, size=10.5, color=BLACK):
    p   = doc.add_paragraph(style='List Bullet')
    run = p.add_run(text)
    run.font.size = Pt(size); run.font.color.rgb = color
    p.paragraph_format.space_after = Pt(3)

def center_text(doc, text, size=9, color=MUTED, italic=True):
    p   = doc.add_paragraph()
    run = p.add_run(text)
    run.font.size = Pt(size); run.font.color.rgb = color; run.italic = italic
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER

def make_table(doc, headers, rows, col_widths=None,
               hdr_bg="1E40AF", row_bg_a="EFF6FF", row_bg_b="FFFFFF"):
    ncols = len(headers)
    t = doc.add_table(rows=1 + len(rows), cols=ncols)
    t.style = 'Table Grid'
    t.alignment = WD_TABLE_ALIGNMENT.LEFT

    hrow = t.rows[0]
    set_row_h(hrow, 0.75)
    for ci, h in enumerate(headers):
        cell = hrow.cells[ci]
        set_cell_bg(cell, hdr_bg)
        p = cell.paragraphs[0]; p.clear()
        r = p.add_run(h)
        r.bold = True; r.font.size = Pt(10.5); r.font.color.rgb = WHITE
        cell.vertical_alignment = WD_ALIGN_VERTICAL.CENTER

    for ri, row_data in enumerate(rows):
        row = t.rows[ri + 1]
        set_row_h(row, 0.72)
        bg = row_bg_a if ri % 2 == 0 else row_bg_b
        for ci, val in enumerate(row_data):
            cell = row.cells[ci]
            set_cell_bg(cell, "DBEAFE" if ci == 0 else bg)
            fc = LBLUE if ci == 0 else BLACK
            p = cell.paragraphs[0]; p.clear()
            r = p.add_run(str(val))
            r.font.size = Pt(10.5); r.font.color.rgb = fc
            r.bold = (ci == 0)
            cell.vertical_alignment = WD_ALIGN_VERTICAL.CENTER

    if col_widths:
        for i, w in enumerate(col_widths):
            for row in t.rows:
                row.cells[i].width = Cm(w)
    return t

# ══════════════════════════════════════════════════════════════
doc = Document()
for sec in doc.sections:
    sec.top_margin = Cm(2.0); sec.bottom_margin = Cm(2.0)
    sec.left_margin = Cm(2.5); sec.right_margin = Cm(2.5)

# ─────────────────────────────────────────────────────────────
# COVER PAGE
# ─────────────────────────────────────────────────────────────
p = doc.add_paragraph()
p.alignment = WD_ALIGN_PARAGRAPH.CENTER
p.paragraph_format.space_before = Pt(52)
r = p.add_run("YouTube Analytics")
r.bold = True; r.font.size = Pt(36); r.font.color.rgb = BLUE

p = doc.add_paragraph()
p.alignment = WD_ALIGN_PARAGRAPH.CENTER
r = p.add_run("Comparative Analysis Report")
r.bold = True; r.font.size = Pt(28); r.font.color.rgb = LBLUE

doc.add_paragraph()
p = doc.add_paragraph()
p.alignment = WD_ALIGN_PARAGRAPH.CENTER
r = p.add_run("NumPy / Pandas Statistical Analysis  vs  Deep Learning (MLP Neural Network)")
r.font.size = Pt(13); r.font.color.rgb = SLATE; r.italic = True

doc.add_paragraph()
p = doc.add_paragraph()
p.alignment = WD_ALIGN_PARAGRAPH.CENTER
r = p.add_run("─────────────────────────────────────")
r.font.color.rgb = BLUE

doc.add_paragraph()

cover_tbl = doc.add_table(rows=5, cols=2)
cover_tbl.alignment = WD_TABLE_ALIGNMENT.CENTER
cover_tbl.style = 'Table Grid'
meta = [
    ("Reports Compared",   "YouTube_Trending_Analysis_Report  ·  YouTube_Views_Prediction_Report"),
    ("Analysis A",         "NumPy / Pandas — Statistical Descriptive Analysis\n10,000 videos · 24 countries · 17 categories"),
    ("Analysis B",         "Deep Learning (MLP) — Time-Series View Forecast\n73 months · 2 independent models"),
    ("Report Date",        datetime.date.today().strftime("%B %d, %Y")),
    ("Classification",     "Analytical Methodology Comparison — Confidential"),
]
for i, (k, v) in enumerate(meta):
    row = cover_tbl.rows[i]; set_row_h(row, 0.90)
    lc = row.cells[0]; rc = row.cells[1]
    set_cell_bg(lc, "1E40AF"); set_cell_bg(rc, "EFF6FF")
    lp = lc.paragraphs[0]; lp.clear()
    lr = lp.add_run(k)
    lr.bold = True; lr.font.size = Pt(10); lr.font.color.rgb = WHITE
    lp.alignment = WD_ALIGN_PARAGRAPH.CENTER
    lc.vertical_alignment = WD_ALIGN_VERTICAL.CENTER
    rp = rc.paragraphs[0]; rp.clear()
    rr = rp.add_run(v)
    rr.font.size = Pt(9.5); rr.font.color.rgb = LBLUE
    rc.vertical_alignment = WD_ALIGN_VERTICAL.CENTER
    lc.width = Cm(4.5); rc.width = Cm(12.0)

doc.add_paragraph()
center_text(doc,
    "Copyright 2024 Meruva Kodanda Suraj  ·  Licensed under Apache License, Version 2.0",
    size=9, color=MUTED)
center_text(doc,
    "http://www.apache.org/licenses/LICENSE-2.0",
    size=9, color=CYAN)

doc.add_page_break()

# ─────────────────────────────────────────────────────────────
# 1. ANALYSIS OVERVIEW
# ─────────────────────────────────────────────────────────────
heading(doc, "1.  Analysis Overview", size=15, color=BLUE)
hr_line(doc)

body(doc,
    "This report compares the analytical outcomes, capabilities, and practical utility of two distinct "
    "data science approaches applied to YouTube data. "
    "Analysis A uses NumPy and Pandas-based descriptive statistics to survey 10,000 trending videos "
    "across 24 countries, extracting broad patterns, correlations, and segment-level benchmarks. "
    "Analysis B applies a Deep Learning MLP neural network to 73 months of single-channel monthly "
    "view counts, learning temporal patterns to forecast views 36 months into the future.")

doc.add_paragraph()
make_table(doc,
    ["Item", "Statistical Analysis (Np · Pd)", "Deep Learning (MLP)"],
    [
        ["Dataset",        "10,000 trending videos · 24 countries",   "Single channel · 73 months (2020–2026)"],
        ["Primary Tools",  "NumPy, Pandas",                           "MLP + Batch Norm + Residual Blocks (Rust)"],
        ["Analysis Type",  "Descriptive · Correlation · Segmentation","Time-series forecasting (36 months)"],
        ["Output",         "Historical & current insights",           "Future predictions (2026–2028)"],
        ["Question Answered", "'What happened and why?'",             "'What will happen next?'"],
        ["Accuracy Metric","CV%, Pearson r",                          "Basic 63% · Advanced 99% (validation set)"],
        ["Execution Time", "Seconds (no training required)",          "Seconds (pre-trained weights loaded)"],
    ],
    col_widths=[4.5, 6.5, 6.5],
)

doc.add_page_break()

# ─────────────────────────────────────────────────────────────
# 2. METHODOLOGY COMPARISON
# ─────────────────────────────────────────────────────────────
heading(doc, "2.  Methodology Comparison", size=15, color=BLUE)
hr_line(doc)

body(doc,
    "The statistical approach is optimised for data exploration and discovery. "
    "Descriptive statistics (mean, median, std dev, CV%), Pearson correlation coefficients, "
    "and cross-tabulations clearly quantify performance differences across every data dimension "
    "— categories, countries, weekdays, months, subscriber tiers, and video formats — "
    "delivering immediately interpretable findings with no model training required.")

body(doc,
    "The deep learning approach specialises in capturing non-linear temporal dependencies. "
    "An MLP with Batch Normalisation and Residual Blocks (Advanced model) or plain dense layers "
    "(Basic model) learns complex seasonality and subscriber growth trends from the historical "
    "time series, projecting future values with log-scale stability. "
    "The first forecast point is anchored to the last observed value to eliminate discontinuity.")

doc.add_paragraph()
make_table(doc,
    ["Methodology Item", "Statistical Analysis (Np · Pd)", "Deep Learning (MLP)"],
    [
        ["Data Preprocessing",   "Pandas filter · groupby · pivot tables",     "JSON load → float scaling → feature vectors"],
        ["Core Algorithm",       "Descriptive stats · Pearson r · CrossTab",    "Linear → BN → GELU → Residual → Head"],
        ["Hyperparameters",      "None (computation-based)",                    "Layer depth · BN epsilon · subscriber slope"],
        ["Training Required",    "No",                                          "Yes — offline, weights serialised to JSON"],
        ["Accuracy Metric",      "CV%, correlation significance",               "Held-out validation accuracy"],
        ["Output Scale",         "Absolute counts (views, likes, engagement)",  "Millions (M), log-scale internally"],
        ["Interpretability",     "Full — every number directly explainable",    "Limited — neural network black-box"],
        ["Reusability",          "Re-run instantly on new data",                "Requires retraining for new channels"],
    ],
    col_widths=[5.0, 6.0, 6.0],
)

doc.add_page_break()

# ─────────────────────────────────────────────────────────────
# 3. ANALYTICAL CAPABILITY COMPARISON
# ─────────────────────────────────────────────────────────────
heading(doc, "3.  Analytical Capability Comparison", size=15, color=BLUE)
hr_line(doc)

body(doc,
    "The analytical strengths of the two approaches are complementary rather than overlapping. "
    "Statistical analysis excels in breadth — it surveys the full landscape of 10,000 videos "
    "and surfaces patterns across dozens of dimensions simultaneously. "
    "Deep learning excels in depth along the time axis — it models complex non-linear trends "
    "that simple statistics cannot capture and produces actionable numeric forecasts.")

doc.add_paragraph()
heading(doc, "  3-1.  Statistical Analysis — Capability Detail", size=12, color=LGREEN, before=8, after=4)

make_table(doc,
    ["Analysis Area", "Key Result", "Capability Rating"],
    [
        ["Descriptive Statistics",
         "Views CV% 1,038% → extreme right-skew confirmed\nMean 3.12M vs Median 59K",
         "★★★★★  Full distribution structure captured"],
        ["Category Benchmarking",
         "Sports 4.65M · Shorts like-rate 6.61%\n17 categories fully decomposed",
         "★★★★★  Complete cross-category comparison"],
        ["Country Analysis",
         "Russia 8.4M vs Philippines 886K (9.5× gap)\n24-country ranking + engagement patterns",
         "★★★★★  Geographic performance quantified"],
        ["Correlation Analysis",
         "All |r| < 0.03 → engagement unpredictable\nfrom surface metrics (views, subscribers, clickbait)",
         "★★★★☆  Null finding is itself a key insight"],
        ["Temporal Patterns",
         "Sunday 5.86M avg · November 6.18M\nQ4 season effect clearly identified",
         "★★★★☆  Weekday & monthly cycles clear"],
        ["Viral Profile",
         "Top 1% = 60.8× overall average views\nEngagement nearly identical to average",
         "★★★★☆  Viral threshold precisely defined"],
        ["Future Forecasting",
         "Not possible — descriptive statistics\nare inherently retrospective",
         "★☆☆☆☆  No predictive capability"],
    ],
    col_widths=[4.0, 7.5, 5.0],
    hdr_bg="065F46",
)

doc.add_paragraph()
heading(doc, "  3-2.  Deep Learning — Capability Detail", size=12, color=LBLUE, before=8, after=4)

make_table(doc,
    ["Analysis Area", "Key Result", "Capability Rating"],
    [
        ["Long-Range Forecasting",
         "36-month monthly predictions at 99% accuracy\n(Advanced model, validation set)",
         "★★★★★  Core and unique strength"],
        ["Seasonality Capture",
         "April peak consistent across 2026, 2027, 2028\nNon-linear seasonal cycles auto-learned",
         "★★★★★  Complex patterns automatically learned"],
        ["Uncertainty Quantification",
         "Model gap defines 0.142M–0.198M range\nTwo-model ensemble provides implicit CI",
         "★★★★☆  Confidence interval expressed via gap"],
        ["Feature Integration",
         "Subscriber slope + seasonal templates\nunified in a single model pipeline",
         "★★★☆☆  Linear subscriber assumption limits accuracy"],
        ["Multi-dimensional Segmentation",
         "Single-channel time-series only\nNo category, country, or format breakdown",
         "★★☆☆☆  Multi-dim segmentation not possible"],
        ["Interpretability",
         "Neural network black-box\nPrediction rationale cannot be directly explained",
         "★★☆☆☆  Low explainability"],
        ["Data Exploration",
         "Focused on one channel's time series\nBroadset pattern discovery not supported",
         "★★☆☆☆  Narrow scope by design"],
    ],
    col_widths=[4.0, 7.5, 5.0],
    hdr_bg="1E40AF",
)

doc.add_page_break()

# ─────────────────────────────────────────────────────────────
# 4. KEY FINDINGS COMPARISON
# ─────────────────────────────────────────────────────────────
heading(doc, "4.  Key Findings Comparison", size=15, color=BLUE)
hr_line(doc)

body(doc,
    "The two analyses illuminate the YouTube ecosystem from fundamentally different vantage points. "
    "Statistical analysis answers 'which content trends and why across the platform', "
    "while deep learning specifically answers 'how will this channel's views evolve over time'. "
    "Together they form a complete picture.")

doc.add_paragraph()
make_table(doc,
    ["Finding Area", "Statistical Analysis (Np · Pd)", "Deep Learning (MLP)"],
    [
        ["Peak Timing",
         "November 6.18M — highest month\nSunday 5.86M — highest day\nQ4 season effect dominant",
         "April is channel peak for 2026–2028\nChannel seasonality differs from platform-wide"],
        ["View Scale",
         "Trending avg 3.12M / median 59K\nViral top 1% = 60.8× overall average",
         "Channel forecast 0.142M–0.198M range\n+57% above 2025 channel avg sustained"],
        ["Growth Trend",
         "2024: +75% spike · 2025: -39% correction\nV-shaped recovery cycle observed",
         "Advanced model: stable plateau\nBasic model: gradual -0.8%/yr decline"],
        ["Engagement Insight",
         "All metrics r < 0.03 vs engagement\nHigh views ≠ high engagement",
         "Engagement not modelled — view count only\nFocused metric: monthly avg views"],
        ["Category Insight",
         "Sports & Shows lead at 4.6M+\nShorts highest like rate at 6.61%",
         "Single channel — category comparison\nnot applicable"],
        ["Country Insight",
         "Russia & Turkey 8M+ (3× the US)\nPhilippines lowest views, highest engagement",
         "No country segmentation — global\nchannel data only"],
        ["Forecasting",
         "None — purely retrospective\nNo future projections possible",
         "36-month monthly forecast\nUpper/lower bounds from two models"],
    ],
    col_widths=[3.5, 7.0, 7.0],
    hdr_bg="1E3A5F",
)

doc.add_page_break()

# ─────────────────────────────────────────────────────────────
# 5. STRENGTHS & WEAKNESSES
# ─────────────────────────────────────────────────────────────
heading(doc, "5.  Strengths & Weaknesses", size=15, color=BLUE)
hr_line(doc)

heading(doc, "  5-1.  Statistical Analysis — Strengths", size=12, color=LGREEN, before=8, after=4)
for b in [
    "Full interpretability — every metric is causally explainable; accessible to non-technical stakeholders",
    "Broad coverage — 10,000 videos, 24 countries, 17 categories analysed simultaneously in one pass",
    "Immediately actionable strategy — weekday, monthly, category, and format insights feed directly into content planning",
    "Zero training cost — full analysis completes in seconds with no model training or GPU required",
    "Outlier and distribution detection — CV% 1,038% instantly reveals the extreme skew in view data",
    "Null-finding value — confirming r < 0.03 between surface metrics and engagement is itself a critical strategic insight",
]: bullet(doc, b, color=GREEN)

doc.add_paragraph()
heading(doc, "  5-2.  Statistical Analysis — Weaknesses", size=12, color=RED, before=8, after=4)
for b in [
    "No future forecasting — descriptive statistics are inherently retrospective; no future projections possible",
    "Correlation ≠ causation — risk of misinterpreting statistical associations as causal mechanisms",
    "Non-linear patterns missed — complex non-linear relationships between variables are not well captured",
    "No temporal continuity — aggregation ignores time ordering and sequential dependencies in the data",
    "Static snapshot — results reflect the dataset period only; shifts in platform behaviour are not projected",
]: bullet(doc, b, color=RED)

doc.add_paragraph()
heading(doc, "  5-3.  Deep Learning — Strengths", size=12, color=LBLUE, before=8, after=4)
for b in [
    "Long-range forecast — projects 36 months of monthly view counts at up to 99% accuracy (Advanced model)",
    "Non-linear learning — seasonality, subscriber growth trends, and momentum shifts are auto-modelled by the MLP",
    "Uncertainty quantification — the prediction gap between two independent models defines a natural confidence interval (0.142M–0.198M)",
    "Feature integration — subscriber growth slope and seasonal templates are unified in a single end-to-end pipeline",
    "Log-scale stability — log-space output handles the wide dynamic range of view counts without numerical issues",
]: bullet(doc, b, color=LBLUE)

doc.add_paragraph()
heading(doc, "  5-4.  Deep Learning — Weaknesses", size=12, color=RED, before=8, after=4)
for b in [
    "Black-box predictions — neural network internals are opaque; the rationale behind any specific forecast is hard to explain",
    "Narrow data scope — single-channel time series only; no cross-category, cross-country, or format-level comparison",
    "External shocks ignored — platform policy changes, competitor actions, and viral events are not incorporated",
    "Linear subscriber assumption — if actual subscriber growth is non-linear, forecast errors compound over the 36-month horizon",
    "Retraining cost — applying the model to a new channel requires offline GPU training and a sufficient historical dataset",
]: bullet(doc, b, color=RED)

doc.add_page_break()

# ─────────────────────────────────────────────────────────────
# 6. OVERALL SCORECARD
# ─────────────────────────────────────────────────────────────
heading(doc, "6.  Overall Scorecard", size=15, color=BLUE)
hr_line(doc)

body(doc,
    "The scorecard below rates both approaches across 8 key analytical dimensions on a 5-point scale. "
    "Stars reflect the capability ceiling of each method for a given dimension, independent of dataset size.")

doc.add_paragraph()
make_table(doc,
    ["Analytical Capability", "Statistical Analysis\n(Np · Pd)", "Deep Learning\n(MLP)", "Winner"],
    [
        ["Data Exploration & Discovery",   "★★★★★", "★★☆☆☆", "Statistical"],
        ["Future Forecasting",             "★☆☆☆☆", "★★★★★", "Deep Learning"],
        ["Interpretability",               "★★★★★", "★★☆☆☆", "Statistical"],
        ["Multi-dimensional Segmentation", "★★★★★", "★★☆☆☆", "Statistical"],
        ["Non-linear Pattern Detection",   "★★☆☆☆", "★★★★★", "Deep Learning"],
        ["Seasonality Capture",            "★★★★☆", "★★★★★", "Deep Learning"],
        ["Execution Speed",                "★★★★★", "★★★☆☆", "Statistical"],
        ["Uncertainty Quantification",     "★★★☆☆", "★★★★☆", "Deep Learning"],
        ["Total Score",                    "26 / 40", "24 / 40", "—"],
    ],
    col_widths=[6.0, 3.5, 3.5, 3.5],
    hdr_bg="1E3A5F",
)

doc.add_paragraph()
body(doc,
    "Overall, the two approaches score comparably (Statistical 26/40 · Deep Learning 24/40), "
    "but the advantage clearly splits by purpose. Statistical analysis dominates for "
    "'understanding the present and past', while deep learning dominates for "
    "'predicting the future'. Neither approach is superior in absolute terms — "
    "the right choice depends entirely on the analytical objective.")

doc.add_page_break()

# ─────────────────────────────────────────────────────────────
# 7. INTEGRATED STRATEGY RECOMMENDATION
# ─────────────────────────────────────────────────────────────
heading(doc, "7.  Integrated Strategy Recommendation", size=15, color=BLUE)
hr_line(doc)

body(doc,
    "The two methodologies are complementary rather than competing. "
    "A two-phase analytical pipeline is recommended for real-world YouTube strategy deployment:")

doc.add_paragraph()
bullet(doc,
    "Phase 1 — Statistical Analysis (NumPy/Pandas): "
    "Explore the full dataset to understand which categories, countries, weekdays, and formats drive performance. "
    "Identify the highest-potential channel segments and key seasonal windows.")
bullet(doc,
    "Phase 2 — Deep Learning (MLP): "
    "Apply the time-series model to channels identified in Phase 1. "
    "Generate 36-month monthly view forecasts with upper and lower bounds for budget and scheduling decisions.")
bullet(doc,
    "Optimal combination: "
    "Statistical analysis answers 'where to focus' — deep learning answers 'how much to expect'. "
    "Combining both delivers a complete, evidence-based YouTube strategy.")

doc.add_paragraph()
make_table(doc,
    ["Scenario", "Recommended Approach", "Key Output"],
    [
        ["Market research before channel launch",
         "Statistical Analysis",
         "Optimal category, country, and upload-day selection"],
        ["Growth forecasting for existing channel",
         "Deep Learning",
         "36-month monthly view range (upper/lower bounds)"],
        ["Content type & format strategy",
         "Statistical Analysis",
         "Clickbait effect, question-title uplift, Shorts ROI"],
        ["Advertising budget planning",
         "Deep Learning",
         "View forecast range → revenue projection → budget ceiling"],
        ["Channel acquisition valuation",
         "Both in combination",
         "Current platform value (stats) + future value (DL)"],
        ["Audience engagement optimisation",
         "Statistical Analysis",
         "Engagement vs surface metrics — where to invest"],
        ["Long-term content calendar",
         "Both in combination",
         "Peak timing from stats · monthly targets from DL"],
    ],
    col_widths=[5.5, 4.5, 6.5],
    hdr_bg="1E40AF",
)

doc.add_page_break()

# ─────────────────────────────────────────────────────────────
# 8. CONCLUSION
# ─────────────────────────────────────────────────────────────
heading(doc, "8.  Conclusion", size=15, color=BLUE)
hr_line(doc)

body(doc,
    "NumPy/Pandas statistical analysis reveals the structural truths of the YouTube ecosystem "
    "across 10,000 trending videos and 24 countries. It confirms that viral success produces a "
    "60× view multiplier, that surface metrics are entirely uncorrelated with engagement (r < 0.03), "
    "that Sunday and November represent the highest-traffic windows, and that Russia and Turkey "
    "outperform the US by a factor of 3× in average views. These findings are immediately "
    "actionable for any content strategist without requiring model training or technical infrastructure.")

body(doc,
    "The Deep Learning MLP models capture temporal patterns invisible to descriptive statistics. "
    "Trained on 73 months of single-channel data, the Advanced model achieves 99% validation accuracy "
    "and projects a stable view plateau of approximately 0.17M per month through 2028 — "
    "more than 57% above the channel's 2025 average. The Basic model's mild-decline forecast "
    "and the Advanced model's stable outlook together define a credible performance corridor "
    "for channel planning.")

body(doc,
    "The fundamental conclusion is that neither approach is universally superior. "
    "Statistical analysis is the tool of choice when the goal is exploration, discovery, "
    "segmentation, and explanation of historical data. Deep learning is the tool of choice "
    "when the goal is prediction, forecasting, and planning for future performance. "
    "Deploying both in a structured two-phase pipeline — 'understand the landscape first, "
    "then forecast the trajectory' — produces the most complete and reliable analytical outcome "
    "for data-driven YouTube channel strategy.")

# ─────────────────────────────────────────────────────────────
# LICENSE NOTICE
# ─────────────────────────────────────────────────────────────
heading(doc, "License", size=13, color=SLATE, before=20, after=4)
hr_line(doc, "9CA3AF")

license_text = (
    "Copyright 2024 Meruva Kodanda Suraj\n\n"
    "Licensed under the Apache License, Version 2.0 (the \"License\");\n"
    "you may not use this file except in compliance with the License.\n"
    "You may obtain a copy of the License at\n\n"
    "    http://www.apache.org/licenses/LICENSE-2.0\n\n"
    "Unless required by applicable law or agreed to in writing, software\n"
    "distributed under the License is distributed on an \"AS IS\" BASIS,\n"
    "WITHOUT WARRANTIES OR CONDITIONS OF ANY KIND, either express or implied.\n"
    "See the License for the specific language governing permissions and\n"
    "limitations under the License."
)
p = doc.add_paragraph()
run = p.add_run(license_text)
run.font.size = Pt(9); run.font.color.rgb = SLATE; run.font.name = "Courier New"
p.paragraph_format.space_after = Pt(12)

# Footer
hr_line(doc, "9CA3AF")
center_text(doc,
    f"Lightgo Analytics  ·  Comparative Analysis Report  ·  "
    f"{datetime.date.today().strftime('%B %Y')}  ·  Confidential",
    size=8.5, color=MUTED)

doc.save(OUT_PATH)
print(f"Report saved:\n  {OUT_PATH}")
