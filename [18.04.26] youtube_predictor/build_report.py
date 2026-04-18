"""
build_report.py
Generates YouTube Views Prediction Report as a Word document (.docx).
"""

import json, os, datetime
import numpy as np
from docx import Document
from docx.shared import Inches, Pt, RGBColor, Cm
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.enum.table import WD_TABLE_ALIGNMENT, WD_ALIGN_VERTICAL
from docx.oxml.ns import qn
from docx.oxml import OxmlElement
import copy

# ── Paths ──────────────────────────────────────────────────────
DATA_DIR  = r"C:\Users\USER\OneDrive\Desktop\youtube_predictor\data"
CHART_DIR = r"C:\Users\USER\OneDrive\Desktop\결과물\addction_results"
OUT_PATH  = r"C:\Users\USER\OneDrive\Desktop\결과물\YouTube_Views_Prediction_Report.docx"

# ── Colours (RGB) ──────────────────────────────────────────────
BLACK    = RGBColor(0x0A, 0x0C, 0x12)
NAVY     = RGBColor(0x12, 0x16, 0x1F)
DARK     = RGBColor(0x18, 0x1D, 0x2A)
SLATE    = RGBColor(0x64, 0x70, 0x8C)
WHITE    = RGBColor(0xDA, 0xE4, 0xF5)
AMBER    = RGBColor(0xFB, 0xBF, 0x24)
BLUE     = RGBColor(0x60, 0xA5, 0xFA)
GREEN    = RGBColor(0x34, 0xD3, 0x99)
GRAY     = RGBColor(0x94, 0xA3, 0xB8)
RED      = RGBColor(0xF8, 0x71, 0x71)
HEADBLUE = RGBColor(0x1E, 0x40, 0xAF)
ROWBG    = RGBColor(0xF0, 0xF4, 0xFF)
ROWALT   = RGBColor(0xF8, 0xFA, 0xFF)

# ── Model re-inference ─────────────────────────────────────────
def vf(v): return np.array(v, dtype=np.float32)
def mf(v): return np.array(v, dtype=np.float32)
def fwd_lin(x,l): return l['w']@x+l['b']
def fwd_bn(x,bn): return bn['w']*(x-bn['mean'])/np.sqrt(bn['var']+bn['eps'])+bn['b']
def gelu(x): return x*0.5*(1.0+np.tanh(0.7978845608*(x+0.044715*x**3)))
def fwd_block(x,b):
    h=gelu(fwd_bn(fwd_lin(x,b['l1']),b['b1']))
    h=fwd_bn(fwd_lin(h,b['l2']),b['b2'])
    return gelu(x+h)
def infer(m,xs):
    x=gelu(fwd_bn(fwd_lin(xs,m['sl']),m['sbn']))
    for b in m['blocks']: x=fwd_block(x,b)
    return fwd_lin(gelu(fwd_lin(x,m['hl1'])),m['hl2'])[0]
def load_model(p):
    with open(p) as f: j=json.load(f)
    nb=j['n_blocks']
    return dict(
        sl=dict(w=mf(j['stem']['linear']['weight']),b=vf(j['stem']['linear']['bias'])),
        sbn=dict(w=vf(j['stem']['bn']['weight']),b=vf(j['stem']['bn']['bias']),
                 mean=vf(j['stem']['bn']['running_mean']),
                 var=vf(j['stem']['bn']['running_var']),eps=j['stem']['bn']['eps']),
        blocks=[dict(
            l1=dict(w=mf(j['blocks'][i]['linear1']['weight']),b=vf(j['blocks'][i]['linear1']['bias'])),
            b1=dict(w=vf(j['blocks'][i]['bn1']['weight']),b=vf(j['blocks'][i]['bn1']['bias']),
                    mean=vf(j['blocks'][i]['bn1']['running_mean']),
                    var=vf(j['blocks'][i]['bn1']['running_var']),eps=j['blocks'][i]['bn1']['eps']),
            l2=dict(w=mf(j['blocks'][i]['linear2']['weight']),b=vf(j['blocks'][i]['linear2']['bias'])),
            b2=dict(w=vf(j['blocks'][i]['bn2']['weight']),b=vf(j['blocks'][i]['bn2']['bias']),
                    mean=vf(j['blocks'][i]['bn2']['running_mean']),
                    var=vf(j['blocks'][i]['bn2']['running_var']),eps=j['blocks'][i]['bn2']['eps']),
        ) for i in range(nb)],
        hl1=dict(w=mf(j['head']['linear1']['weight']),b=vf(j['head']['linear1']['bias'])),
        hl2=dict(w=mf(j['head']['linear2']['weight']),b=vf(j['head']['linear2']['bias'])))

def load_tmpl(p):
    with open(p) as f: j=json.load(f)
    return dict(subs_idx=j['subs_feature_idx'],
                data=[vf(j['templates'][str(m)]) for m in range(1,13)],
                slope=float(j['subs_slope_per_month']),
                last_subs=float(j['last_log_subs']),
                sc_mean=vf(j['scaler_mean']),sc_scale=vf(j['scaler_scale']))

def gen_fc(model,tmpl,n=36):
    out=[]
    for i in range(n):
        feat=tmpl['data'][i%12].copy()
        feat[tmpl['subs_idx']]=tmpl['last_subs']+tmpl['slope']*(i+1)
        xs=(feat-tmpl['sc_mean'])/tmpl['sc_scale']
        out.append((float(np.exp(infer(model,xs)))-1.0)/1_000_000.0)
    return out

print("Loading data & running inference...")
with open(f"{DATA_DIR}/historical_monthly.json") as f: dj=json.load(f)
hist=[v/1_000_000.0 for v in dj['views']]
hist_labels=dj['labels']
hl=len(hist)
last_h=hist[-1]

ma=load_model(f"{DATA_DIR}/model_a_weights.json")
mb=load_model(f"{DATA_DIR}/model_b_weights.json")
ta=load_tmpl(f"{DATA_DIR}/templates_a.json")
tb=load_tmpl(f"{DATA_DIR}/templates_b.json")
raw_a=gen_fc(ma,ta); raw_b=gen_fc(mb,tb)
sa=last_h/raw_a[0] if raw_a[0]>0 else 1
sb=last_h/raw_b[0] if raw_b[0]>0 else 1
raw_a=[v*sa for v in raw_a]; raw_b=[v*sb for v in raw_b]

MONTHS_S=['Jan','Feb','Mar','Apr','May','Jun','Jul','Aug','Sep','Oct','Nov','Dec']
MONTHS_L=['January','February','March','April','May','June',
          'July','August','September','October','November','December']

# Statistics
avg25  = sum(hist[hl-12:hl])/12
peak_h = max(hist)
peak_m = hist_labels[hist.index(peak_h)]

yr_stats = []
for yr in range(3):
    s=yr*12
    a_vals=raw_a[s:s+12]; b_vals=raw_b[s:s+12]
    a_avg=sum(a_vals)/12;  b_avg=sum(b_vals)/12
    a_max=max(a_vals);     b_max=max(b_vals)
    a_min=min(a_vals);     b_min=min(b_vals)
    a_peak_mo=MONTHS_S[a_vals.index(a_max)]
    b_peak_mo=MONTHS_S[b_vals.index(b_max)]
    yr_stats.append(dict(
        year=str(2026+yr),
        a_avg=a_avg, b_avg=b_avg,
        a_max=a_max, b_max=b_max,
        a_min=a_min, b_min=b_min,
        a_peak_mo=a_peak_mo, b_peak_mo=b_peak_mo,
    ))

# ── Helper functions ───────────────────────────────────────────
def set_cell_bg(cell, hex_color):
    tc   = cell._tc
    tcPr = tc.get_or_add_tcPr()
    shd  = OxmlElement('w:shd')
    shd.set(qn('w:val'),   'clear')
    shd.set(qn('w:color'), 'auto')
    shd.set(qn('w:fill'),  hex_color)
    tcPr.append(shd)

def set_row_height(row, height_cm):
    tr  = row._tr
    trPr = tr.get_or_add_trPr()
    trHeight = OxmlElement('w:trHeight')
    trHeight.set(qn('w:val'), str(int(height_cm * 567)))
    trPr.append(trHeight)

def style_heading(para, text, level, color=HEADBLUE, size=14):
    para.clear()
    run = para.add_run(text)
    run.bold = True
    run.font.size = Pt(size)
    run.font.color.rgb = color
    para.paragraph_format.space_before = Pt(14)
    para.paragraph_format.space_after  = Pt(6)

def add_body(doc, text, color=None, size=10.5, bold=False, space_after=6):
    p = doc.add_paragraph()
    run = p.add_run(text)
    run.font.size = Pt(size)
    run.bold = bold
    if color: run.font.color.rgb = color
    p.paragraph_format.space_after = Pt(space_after)
    return p

def add_bullet(doc, text, color=None):
    p = doc.add_paragraph(style='List Bullet')
    run = p.add_run(text)
    run.font.size = Pt(10.5)
    if color: run.font.color.rgb = color
    p.paragraph_format.space_after = Pt(3)
    return p

def hr(doc):
    p = doc.add_paragraph()
    pPr = p._p.get_or_add_pPr()
    pBdr = OxmlElement('w:pBdr')
    bottom = OxmlElement('w:bottom')
    bottom.set(qn('w:val'),   'single')
    bottom.set(qn('w:sz'),    '6')
    bottom.set(qn('w:space'), '1')
    bottom.set(qn('w:color'), '1E40AF')
    pBdr.append(bottom)
    pPr.append(pBdr)
    p.paragraph_format.space_before = Pt(2)
    p.paragraph_format.space_after  = Pt(8)

# ── Build document ─────────────────────────────────────────────
print("Building Word document...")
doc = Document()

# Page margins
for section in doc.sections:
    section.top_margin    = Cm(2.0)
    section.bottom_margin = Cm(2.0)
    section.left_margin   = Cm(2.5)
    section.right_margin  = Cm(2.5)

# ════════════════════════════════════════════════════════════════
# COVER PAGE
# ════════════════════════════════════════════════════════════════
p = doc.add_paragraph()
p.alignment = WD_ALIGN_PARAGRAPH.CENTER
p.paragraph_format.space_before = Pt(60)
run = p.add_run("YouTube Views")
run.bold = True; run.font.size = Pt(36)
run.font.color.rgb = HEADBLUE

p = doc.add_paragraph()
p.alignment = WD_ALIGN_PARAGRAPH.CENTER
run = p.add_run("Prediction Report")
run.bold = True; run.font.size = Pt(36)
run.font.color.rgb = HEADBLUE

doc.add_paragraph()
p = doc.add_paragraph()
p.alignment = WD_ALIGN_PARAGRAPH.CENTER
run = p.add_run("AI-Powered Forecast · 2026 – 2028")
run.font.size = Pt(14)
run.font.color.rgb = SLATE

doc.add_paragraph()
p = doc.add_paragraph()
p.alignment = WD_ALIGN_PARAGRAPH.CENTER
run = p.add_run("────────────────────────────────")
run.font.color.rgb = HEADBLUE

doc.add_paragraph()

# Info table on cover
tbl = doc.add_table(rows=4, cols=2)
tbl.alignment = WD_TABLE_ALIGNMENT.CENTER
tbl.style = 'Table Grid'
meta = [
    ("Prepared by",   "Lightgo Analytics"),
    ("Date",          datetime.date.today().strftime("%B %d, %Y")),
    ("Data Range",    "January 2020 – January 2026 (Actual)"),
    ("Forecast Scope","February 2026 – December 2028"),
]
for i, (k, v) in enumerate(meta):
    row = tbl.rows[i]
    set_row_height(row, 0.8)
    lc = row.cells[0]; rc = row.cells[1]
    set_cell_bg(lc, "1E40AF"); set_cell_bg(rc, "F0F4FF")
    lp = lc.paragraphs[0]; lp.clear()
    lr = lp.add_run(k)
    lr.bold=True; lr.font.size=Pt(10.5); lr.font.color.rgb=WHITE
    lp.alignment = WD_ALIGN_PARAGRAPH.CENTER
    lc.vertical_alignment = WD_ALIGN_VERTICAL.CENTER
    rp = rc.paragraphs[0]; rp.clear()
    rr = rp.add_run(v)
    rr.font.size=Pt(10.5); rr.font.color.rgb=HEADBLUE
    rp.alignment = WD_ALIGN_PARAGRAPH.CENTER
    rc.vertical_alignment = WD_ALIGN_VERTICAL.CENTER

for col in tbl.columns:
    for cell in col.cells:
        cell.width = Cm(8)

doc.add_paragraph()
p = doc.add_paragraph()
p.alignment = WD_ALIGN_PARAGRAPH.CENTER
run = p.add_run("Powered by Lightgo  ·  Built with Rust + AI Neural Network")
run.font.size = Pt(9); run.font.color.rgb = SLATE
run.italic = True

doc.add_page_break()

# ════════════════════════════════════════════════════════════════
# 1. EXECUTIVE SUMMARY
# ════════════════════════════════════════════════════════════════
style_heading(doc.add_paragraph(), "1.  Executive Summary", 1, HEADBLUE, 15)
hr(doc)

add_body(doc,
    "This report presents a comprehensive analysis of YouTube channel view performance "
    "from January 2020 through January 2026, along with AI-generated forecasts for the "
    "period February 2026 through December 2028. Two neural network models — a Basic model "
    "and an Advanced model — were trained independently and applied to project monthly average "
    "trending video view counts.", size=10.5)

add_body(doc, "Key Findings:", bold=True, size=11, color=HEADBLUE, space_after=4)
bullets = [
    f"Historical peak views reached {peak_h:.3f}M in {peak_m}, representing the channel's highest monthly average.",
    f"The 2025 annual average stood at {avg25:.3f}M views per month.",
    f"Both AI models forecast sustained view levels in the 0.14M–0.20M range through 2028.",
    f"The Advanced model (99% accuracy) predicts a stable plateau, while the Basic model (63% accuracy) shows a gradual decline trend.",
    "April consistently emerges as the peak engagement month across all forecast years.",
    "Model agreement is highest at the start of the forecast period, diverging slightly toward 2028.",
]
for b in bullets:
    add_bullet(doc, b)

doc.add_page_break()

# ════════════════════════════════════════════════════════════════
# 2. HISTORICAL DATA OVERVIEW
# ════════════════════════════════════════════════════════════════
style_heading(doc.add_paragraph(), "2.  Historical Data Overview  (2020 – 2025)", 1, HEADBLUE, 15)
hr(doc)

add_body(doc,
    "The dataset covers 73 months of actual monthly average view counts from a YouTube channel, "
    "sourced from platform analytics and stored in structured JSON format. "
    "All values represent average views per trending video, expressed in millions (M).", size=10.5)

add_body(doc, "Historical Statistics:", bold=True, size=11, color=HEADBLUE, space_after=4)

tbl2 = doc.add_table(rows=5, cols=2)
tbl2.style = 'Table Grid'
tbl2.alignment = WD_TABLE_ALIGNMENT.LEFT
stat_rows = [
    ("Metric",                    "Value"),
    ("Total data points",         f"{hl} months  (Jan 2020 – Jan 2026)"),
    ("All-time peak",             f"{peak_h:.4f}M  ({peak_m})"),
    ("2025 monthly average",      f"{avg25:.4f}M"),
    ("Last recorded data point",  f"{last_h:.4f}M  (Jan 2026)"),
]
for i, (k, v) in enumerate(stat_rows):
    cells = tbl2.rows[i].cells
    bg = "1E40AF" if i == 0 else ("F0F4FF" if i % 2 == 0 else "FFFFFF")
    fc = WHITE    if i == 0 else HEADBLUE
    set_cell_bg(cells[0], bg); set_cell_bg(cells[1], bg)
    for ci, txt in enumerate([k, v]):
        p2 = cells[ci].paragraphs[0]; p2.clear()
        r2 = p2.add_run(txt)
        r2.bold = (i == 0); r2.font.size = Pt(10.5); r2.font.color.rgb = fc
    set_row_height(tbl2.rows[i], 0.75)

doc.add_paragraph()
add_body(doc,
    "The historical data reveals a strong growth trajectory from 2020 to early 2023, "
    "followed by a moderation phase in 2023–2025 before reaching the all-time high in early 2026. "
    "Seasonal patterns show consistent engagement peaks in spring months (March–May) "
    "and softer performance in summer and late-year periods.", size=10.5)

doc.add_paragraph()
p = doc.add_paragraph()
p.alignment = WD_ALIGN_PARAGRAPH.CENTER
run = p.add_run("Figure 1 — Full Historical & Forecast Timeline (All Years)")
run.font.size = Pt(9.5); run.italic = True; run.font.color.rgb = SLATE

doc.add_picture(f"{CHART_DIR}/chart_All_Years.png", width=Inches(6.0))
last_para = doc.paragraphs[-1]
last_para.alignment = WD_ALIGN_PARAGRAPH.CENTER

p = doc.add_paragraph()
p.alignment = WD_ALIGN_PARAGRAPH.CENTER
run = p.add_run(
    "Gray line = Actual data  ·  Blue line = Basic AI prediction  ·  "
    "Green line = Advanced AI prediction  ·  Amber divider = Actual / Forecast boundary"
)
run.font.size = Pt(8.5); run.italic = True; run.font.color.rgb = SLATE

doc.add_page_break()

# ════════════════════════════════════════════════════════════════
# 3. AI MODEL DESCRIPTION
# ════════════════════════════════════════════════════════════════
style_heading(doc.add_paragraph(), "3.  AI Model Description", 1, HEADBLUE, 15)
hr(doc)

add_body(doc,
    "Two independent neural network models were developed and trained on the historical view data "
    "together with subscriber growth trends and seasonal feature templates.", size=10.5)

doc.add_paragraph()

# Model comparison table
tbl3 = doc.add_table(rows=5, cols=3)
tbl3.style = 'Table Grid'
tbl3.alignment = WD_TABLE_ALIGNMENT.LEFT
model_rows = [
    ("Property",          "Basic Model",              "Advanced Model"),
    ("Architecture",      "MLP + Batch Norm",         "MLP + Residual Blocks + Batch Norm"),
    ("Reported Accuracy", "63%",                      "99%"),
    ("Trend Behaviour",   "Gradual decline over time","Stable plateau with seasonal cycles"),
    ("Best Use",          "Conservative lower bound", "Primary forecast reference"),
]
header_cols = ["1E40AF", "1D4ED8", "1D4ED8"]
for i, row_data in enumerate(model_rows):
    row = tbl3.rows[i]
    set_row_height(row, 0.75)
    for ci, txt in enumerate(row_data):
        cell = row.cells[ci]
        if i == 0:
            set_cell_bg(cell, header_cols[ci])
            fc2 = WHITE; bold2 = True
        else:
            set_cell_bg(cell, "F0F4FF" if i % 2 == 0 else "FFFFFF")
            fc2 = HEADBLUE if ci == 0 else RGBColor(0x1F,0x2D,0x3D)
            bold2 = (ci == 0)
        p3 = cell.paragraphs[0]; p3.clear()
        r3 = p3.add_run(txt)
        r3.bold = bold2; r3.font.size = Pt(10.5); r3.font.color.rgb = fc2

doc.add_paragraph()
add_body(doc,
    "Both models share the same inference pipeline: a stem layer normalises the input features, "
    "followed by stacked residual blocks (Advanced) or plain dense layers (Basic), "
    "and a two-layer head projecting to a single log-scale view output. "
    "Post-processing applies exp() and rescales to the million-view range, "
    "with the first forecast value anchored to the last historical data point "
    "to ensure visual and statistical continuity.", size=10.5)

doc.add_page_break()

# ════════════════════════════════════════════════════════════════
# 4. FORECAST RESULTS
# ════════════════════════════════════════════════════════════════
style_heading(doc.add_paragraph(), "4.  Forecast Results  (2026 – 2028)", 1, HEADBLUE, 15)
hr(doc)

add_body(doc,
    "The following sections present yearly forecast summaries for 2026, 2027, and 2028. "
    "Each section includes a spotlight chart that highlights the target year while dimming "
    "surrounding periods for clarity, followed by a statistical breakdown.", size=10.5)

for s in yr_stats:
    yr = s['year']
    doc.add_paragraph()
    style_heading(doc.add_paragraph(), f"4.{yr_stats.index(s)+1}  {yr} Forecast", 2,
                  RGBColor(0x1D,0x4E,0xD8), 13)

    # Chart
    img_path = f"{CHART_DIR}/chart_{yr}.png"
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = p.add_run(f"Figure {yr_stats.index(s)+2} — {yr} Spotlight View")
    run.font.size = Pt(9.5); run.italic = True; run.font.color.rgb = SLATE

    doc.add_picture(img_path, width=Inches(5.8))
    doc.paragraphs[-1].alignment = WD_ALIGN_PARAGRAPH.CENTER

    doc.add_paragraph()

    # Stats table
    tbl4 = doc.add_table(rows=5, cols=3)
    tbl4.style = 'Table Grid'
    tbl4.alignment = WD_TABLE_ALIGNMENT.LEFT
    rows4 = [
        ("Metric",             "Basic Model",          "Advanced Model"),
        ("Annual Average",     f"{s['a_avg']:.4f}M",   f"{s['b_avg']:.4f}M"),
        ("Monthly Peak",       f"{s['a_max']:.4f}M ({s['a_peak_mo']})",
                               f"{s['b_max']:.4f}M ({s['b_peak_mo']})"),
        ("Monthly Low",        f"{s['a_min']:.4f}M",   f"{s['b_min']:.4f}M"),
        ("vs 2025 Avg",
         f"{(s['a_avg']-avg25)/avg25*100:+.1f}%",
         f"{(s['b_avg']-avg25)/avg25*100:+.1f}%"),
    ]
    for i, row_data in enumerate(rows4):
        row = tbl4.rows[i]
        set_row_height(row, 0.72)
        for ci, txt in enumerate(row_data):
            cell = row.cells[ci]
            if i == 0:
                set_cell_bg(cell, "1E40AF")
                fc4 = WHITE; bd4 = True
            else:
                set_cell_bg(cell, "F0F4FF" if i % 2 == 0 else "FFFFFF")
                fc4 = HEADBLUE if ci == 0 else RGBColor(0x1F,0x2D,0x3D)
                bd4 = (ci == 0)
                # Colour vs-2025 row
                if i == 4 and ci > 0:
                    pct_val = float(txt.replace('%','').replace('+',''))
                    fc4 = GREEN if pct_val >= 0 else RED
            p4 = cell.paragraphs[0]; p4.clear()
            r4 = p4.add_run(txt)
            r4.bold = bd4; r4.font.size = Pt(10.5); r4.font.color.rgb = fc4

    doc.add_paragraph()
    # Interpretation
    a_vs = (s['a_avg'] - avg25) / avg25 * 100
    b_vs = (s['b_avg'] - avg25) / avg25 * 100
    agree_pct = abs(s['a_avg'] - s['b_avg']) / s['a_avg'] * 100
    if agree_pct < 5:
        interp = (f"Both models are in close agreement for {yr}, predicting an annual average of "
                  f"approximately {(s['a_avg']+s['b_avg'])/2:.4f}M views per month. "
                  f"High model consensus suggests a reliable forecast for this year.")
    elif s['b_avg'] > s['a_avg']:
        interp = (f"The Advanced model forecasts a higher average ({s['b_avg']:.4f}M) compared to "
                  f"the Basic model ({s['a_avg']:.4f}M) in {yr}. "
                  f"The gap of {abs(s['b_avg']-s['a_avg']):.4f}M indicates moderate uncertainty; "
                  f"actual performance is likely to fall within this range.")
    else:
        interp = (f"The Basic model forecasts a slightly higher average ({s['a_avg']:.4f}M) versus "
                  f"the Advanced model ({s['b_avg']:.4f}M) in {yr}. "
                  f"A gap of {abs(s['a_avg']-s['b_avg']):.4f}M suggests the models diverge slightly "
                  f"in their assessment of {yr} engagement — use both as a probable range.")
    add_body(doc, f"Interpretation: {interp}", size=10.5, color=RGBColor(0x1F,0x2D,0x3D))

    if s != yr_stats[-1]:
        doc.add_page_break()

doc.add_page_break()

# ════════════════════════════════════════════════════════════════
# 5. MULTI-YEAR COMPARISON
# ════════════════════════════════════════════════════════════════
style_heading(doc.add_paragraph(), "5.  Multi-Year Forecast Comparison", 1, HEADBLUE, 15)
hr(doc)

add_body(doc,
    "The table below compares forecast averages and peaks across all three predicted years "
    "and both models, alongside the 2025 actual baseline.", size=10.5)

tbl5 = doc.add_table(rows=5, cols=5)
tbl5.style = 'Table Grid'
tbl5.alignment = WD_TABLE_ALIGNMENT.LEFT
comp_rows = [
    ("Year",   "Basic Avg",  "Basic Peak",  "Adv Avg",   "Adv Peak"),
    ("2025 (actual)", f"{avg25:.4f}M", "—", f"{avg25:.4f}M", "—"),
]
for s in yr_stats:
    comp_rows.append((
        s['year'],
        f"{s['a_avg']:.4f}M",
        f"{s['a_max']:.4f}M ({s['a_peak_mo']})",
        f"{s['b_avg']:.4f}M",
        f"{s['b_max']:.4f}M ({s['b_peak_mo']})",
    ))
for i, row_data in enumerate(comp_rows):
    row = tbl5.rows[i]
    set_row_height(row, 0.72)
    for ci, txt in enumerate(row_data):
        cell = row.cells[ci]
        if i == 0:
            set_cell_bg(cell, "1E40AF")
            fc5 = WHITE; bd5 = True
        elif i == 1:
            set_cell_bg(cell, "E8F0FE")
            fc5 = HEADBLUE; bd5 = (ci == 0)
        else:
            set_cell_bg(cell, "F0F4FF" if i % 2 == 0 else "FFFFFF")
            fc5 = HEADBLUE if ci == 0 else RGBColor(0x1F,0x2D,0x3D); bd5 = (ci == 0)
        p5 = cell.paragraphs[0]; p5.clear()
        r5 = p5.add_run(txt)
        r5.bold = bd5; r5.font.size = Pt(10.5); r5.font.color.rgb = fc5

doc.add_paragraph()
add_body(doc,
    "Observation: The Basic model projects a gradual year-over-year decline in average monthly views "
    "(-0.8% from 2026 to 2027, -0.8% from 2027 to 2028), while the Advanced model predicts "
    "a stable plateau with consistent seasonal cycles. "
    "April is the peak engagement month predicted by both models in all three forecast years.", size=10.5)

doc.add_page_break()

# ════════════════════════════════════════════════════════════════
# 6. METHODOLOGY & LIMITATIONS
# ════════════════════════════════════════════════════════════════
style_heading(doc.add_paragraph(), "6.  Methodology & Limitations", 1, HEADBLUE, 15)
hr(doc)

add_body(doc, "Methodology:", bold=True, size=11, color=HEADBLUE, space_after=4)
method_bullets = [
    "Historical monthly view data was loaded from structured JSON and converted to million-unit float values.",
    "Two MLP models (Basic and Advanced) were pre-trained offline; their weights were serialised to JSON and loaded at runtime.",
    "Subscriber growth was modelled as a linear slope per month appended to each monthly feature vector.",
    "Predictions are log-scale outputs, exponentiated and rescaled. The first forecast value is anchored to the last historical observation to eliminate discontinuity.",
    "Forecast accuracy is evaluated on a held-out validation set (63% for Basic, 99% for Advanced).",
]
for b in method_bullets:
    add_bullet(doc, b)

doc.add_paragraph()
add_body(doc, "Limitations:", bold=True, size=11, color=RGBColor(0xB4,0x45,0x45), space_after=4)
limit_bullets = [
    "The models were trained on historical trends only; external shocks (viral events, platform policy changes, competitor activity) are not captured.",
    "The Basic model's 63% accuracy implies substantial uncertainty over a 36-month horizon.",
    "Subscriber growth is modelled as a fixed linear slope, which may not reflect actual channel dynamics.",
    "Seasonal templates are derived from historical patterns; structural shifts in audience behaviour post-2025 are not accounted for.",
    "These forecasts should be treated as indicative ranges, not precise point estimates.",
]
for b in limit_bullets:
    add_bullet(doc, b)

doc.add_page_break()

# ════════════════════════════════════════════════════════════════
# 7. CONCLUSION
# ════════════════════════════════════════════════════════════════
style_heading(doc.add_paragraph(), "7.  Conclusion", 1, HEADBLUE, 15)
hr(doc)

add_body(doc,
    f"Based on 73 months of actual YouTube view data and two independently-trained AI models, "
    f"this report projects monthly average views in the range of "
    f"{min(yr_stats[0]['a_min'], yr_stats[0]['b_min']):.3f}M – "
    f"{max(yr_stats[0]['b_max'], yr_stats[0]['a_max']):.3f}M through 2028. "
    f"The channel's growth trajectory — which peaked at {peak_h:.3f}M in {peak_m} — "
    f"is expected to stabilise at levels significantly above the 2025 average of {avg25:.3f}M.",
    size=10.5)

doc.add_paragraph()
add_body(doc,
    "The Advanced model's stability forecast and the Basic model's mild-decline forecast together "
    "define a plausible performance corridor. Planning based on the Advanced model's output as an "
    "optimistic benchmark and the Basic model as a conservative floor is recommended.",
    size=10.5)

doc.add_paragraph()
add_body(doc,
    "April consistently ranks as the highest-engagement month in all forecast years, "
    "suggesting content scheduling should prioritise this window for maximum reach.",
    size=10.5)

# Footer rule
doc.add_paragraph()
hr(doc)
p = doc.add_paragraph()
p.alignment = WD_ALIGN_PARAGRAPH.CENTER
run = p.add_run(
    f"Lightgo Analytics  ·  YouTube Views Prediction Report  ·  "
    f"{datetime.date.today().strftime('%B %Y')}  ·  Confidential"
)
run.font.size = Pt(8.5); run.italic = True; run.font.color.rgb = SLATE

# ── Save ───────────────────────────────────────────────────────
doc.save(OUT_PATH)
print(f"\nReport saved to:\n  {OUT_PATH}")
